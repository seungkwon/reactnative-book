import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "06_검색으로_목록_필터링하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-filter-list"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "06_검색으로_목록_필터링하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "6장. 검색으로 목록 필터링하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "검색어 입력에 따라 FlatList 결과를 좁혀 보기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "검색어 상태를 만들고, 원본 배열을 조건에 맞게 필터링해 FlatList 결과를 실시간으로 바꾸는 방법을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 로컬 배열 검색에 집중한다. 서버 API 검색이나 디바운스 최적화는 이후 단계에서 다룰 수 있도록 남겨 둔다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch6_filter_list.jpg"
    annotated_app_js = "\n".join(
        [
            "import { useState } from 'react';",
            "import { FlatList, TextInput } from 'react-native';  ①",
            "",
            "export default function App() {",
            "  const [profiles, setProfiles] = useState(initialProfiles);",
            "  const [query, setQuery] = useState('');  ②",
            "",
            "  const normalizedQuery = query.trim().toLowerCase();",
            "  const filteredProfiles = profiles.filter((profile) => {  ③",
            "    if (!normalizedQuery) {",
            "      return true;",
            "    }",
            "",
            "    return (",
            "      profile.name.toLowerCase().includes(normalizedQuery) ||  ④",
            "      profile.role.toLowerCase().includes(normalizedQuery)",
            "    );",
            "  });",
            "",
            "  return (",
            "    <>",
            "      <TextInput value={query} onChangeText={setQuery} />  ⑤",
            "      <FlatList data={filteredProfiles} />  ⑥",
            "    </>",
            "  );",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "6장. 검색으로 목록 필터링하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "6.1 왜 목록 필터링이 필요한가", 1)
    add_body(document, "목록 데이터가 조금만 많아져도 사용자는 원하는 항목을 바로 찾기 어려워진다. 그래서 실제 앱에서는 검색어나 조건에 따라 목록을 좁혀 보여 주는 기능이 자주 필요하다.")
    add_body(document, "이번 장에서는 chapter 05까지 다뤘던 목록 구조를 그대로 유지하면서, 입력창의 검색어에 따라 FlatList 결과가 즉시 달라지는 흐름을 구현한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 검색어 상태와 로컬 배열 필터링")

    add_heading(document, "6.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 추가 라이브러리 없이 useState와 JavaScript 배열 메서드 filter만으로도 충분히 실용적인 검색 UI를 만들 수 있다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "검색 기능을 처음 만들 때는 원본 배열을 직접 바꾸기보다, 검색어를 기준으로 보여 줄 결과만 따로 계산하는 방식이 더 안전하고 이해하기 쉽다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "6.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 06 예제는 이름이나 역할에 검색어가 포함된 프로필만 화면에 보여 준다. 예를 들어 `react`라고 입력하면 역할 텍스트에 해당 단어가 들어간 카드만 남게 된다.")
    add_command_block(
        document,
        [
            "cd chapters/06_검색으로_목록_필터링하기/examples/expo-filter-list",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 검색어를 입력해 보자. 검색어가 바뀔 때마다 결과 목록이 즉시 줄어들면 상태와 필터 계산이 제대로 연결된 것이다.")

    add_heading(document, "6.4 검색어로 결과 계산하기", 2)
    add_body(document, "이번 예제의 핵심은 원본 목록을 훼손하지 않고, 검색어를 기준으로 보여 줄 결과만 따로 계산하는 것이다. 이렇게 하면 검색어를 지웠을 때 원래 목록을 쉽게 다시 보여 줄 수 있다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "이번 장의 핵심 도구는 TextInput과 FlatList다. 하나는 검색어를 받고, 다른 하나는 필터링된 결과를 화면에 표시한다."),
            ("②", "query 상태는 사용자가 입력 중인 검색어를 저장한다. 목록 상태와 별도로 관리해야 검색어 변경과 데이터 변경을 분리해 이해할 수 있다."),
            ("③", "filteredProfiles는 원본 profiles에서 조건에 맞는 항목만 골라낸 결과다. 원본 배열을 직접 수정하지 않고 새로운 결과를 계산하는 방식이 안전하다."),
            ("④", "이름과 역할 모두를 검색 대상으로 두면 사용성이 좋아진다. 사용자는 정확한 이름을 몰라도 관련 키워드로 항목을 찾을 수 있다."),
            ("⑤", "TextInput의 값이 query 상태와 연결돼 있기 때문에, 한 글자씩 입력하는 순간마다 화면 결과가 다시 계산된다."),
            ("⑥", "FlatList는 더 이상 전체 profiles가 아니라 filteredProfiles를 받는다. 바로 이 연결 덕분에 검색어가 곧 화면 결과 변화로 이어진다."),
        ],
    )
    add_body(document, "검색 기능에서 중요한 것은 상태를 많이 만드는 것이 아니라, 어떤 값을 원본으로 두고 어떤 값을 화면용 계산 결과로 둘지 역할을 분명히 나누는 것이다.")

    document.add_page_break()

    add_heading(document, "6.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 검색어로 `react`를 입력한 예시다. 역할 텍스트에 해당 키워드가 포함된 카드만 남아, 사용자가 원하는 항목을 훨씬 빠르게 찾을 수 있다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 6-1. 검색어로 필터링한 FlatList 결과 화면", width_cm=7.6)

    add_heading(document, "6.6 정리", 2)
    add_body(document, "이 장에서는 검색어 상태를 바탕으로 FlatList 결과를 실시간으로 필터링했다. 독자는 목록 데이터 전체를 바꾸지 않고도, 보여 주는 결과만 유연하게 제어할 수 있다는 점을 확인할 수 있다.")
    add_body(document, "다음 장에서는 검색 결과를 정렬하거나, 여러 조건을 함께 적용하는 방식으로 확장하면 목록 처리 능력이 한층 더 실전적으로 발전한다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 검색어를 입력할 때 목록 결과가 즉시 줄어드는지, 그리고 검색어를 지우면 전체 목록이 다시 나타나는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
