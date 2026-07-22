import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "11_WebSocket으로_실시간_채팅_화면_만들기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-chat-ui"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "11_WebSocket으로_실시간_채팅_화면_만들기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "11장 WebSocket으로 실시간 채팅 화면 만들기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "메시지 목록, 입력창, 연결 상태를 한 화면에 담아 채팅 UI의 기본 흐름 익히기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "React Native에서 WebSocket 연결 상태를 관리하고, 메시지 목록과 입력창을 연결해 실시간 채팅 화면의 기본 구조를 완성한다. 12장에서 서버를 붙이기 전에 클라이언트 흐름을 먼저 이해하는 것이 목표다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "이번 장은 채팅 서버를 직접 구현하지 않고, 클라이언트 화면과 상태 흐름에 집중한다. WebSocket 이벤트를 어디에서 받고, 어떤 식으로 화면 상태에 반영하는지 이해해 두면 다음 장에서 서버를 연결할 때 훨씬 수월하다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch11_websocket_chat.jpg"
    annotated_app_js = "\n".join(
        [
            "const WS_URL = 'ws://192.168.0.10:8080'; ①",
            "const USER_ID = 'me';",
            "",
            "export default function App() {",
            "  const socketRef = useRef(null); ②",
            "  const [connectionStatus, setConnectionStatus] = useState('연결 준비 중');",
            "  const [messages, setMessages] = useState(initialMessages);",
            "",
            "  useEffect(() => {",
            "    const socket = new WebSocket(WS_URL); ③",
            "    socketRef.current = socket;",
            "",
            "    socket.onopen = () => {",
            "      setConnectionStatus('연결됨');",
            "    };",
            "",
            "    socket.onmessage = (event) => { ④",
            "      setMessages((current) => [",
            "        ...current,",
            "        { id: `remote-${Date.now()}`, sender: 'other', text: event.data, time: formatTime() },",
            "      ]);",
            "    };",
            "  }, []);",
            "",
            "  const sendMessage = () => {",
            "    const trimmed = inputText.trim();",
            "    setMessages((current) => [...current, { id: `local-${Date.now()}`, sender: USER_ID, text: trimmed, time: formatTime() }]); ⑤",
            "",
            "    if (socketRef.current?.readyState === WebSocket.OPEN) {",
            "      socketRef.current.send(trimmed); ⑥",
            "    }",
            "  };",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "11장 WebSocket으로 실시간 채팅 화면 만들기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "11.1 왜 채팅 UI를 먼저 만들어야 할까", 1)
    add_body(document, "실시간 채팅 기능을 만들 때 많은 입문자가 서버 구현부터 떠올리지만, 실제 앱 개발에서는 먼저 화면 구조와 상태 흐름을 정리해 두는 편이 훨씬 안정적이다. 어떤 메시지를 어떤 형태로 보여줄지, 연결 상태를 어디에 표시할지, 전송 버튼을 눌렀을 때 화면이 어떻게 반응해야 하는지를 먼저 정리하면 서버와 연결할 때도 기준이 분명해진다.")
    add_body(document, "이번 장에서는 WebSocket 서버가 아직 완성되지 않았더라도 채팅 화면 자체는 충분히 만들 수 있다는 점을 경험해 본다. 메시지 배열, 연결 상태 문자열, 입력창 상태만 잘 관리해도 실시간 채팅 앱의 핵심 구조가 이미 갖춰진다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터 또는 실제 Android 기기")
    add_bullet(document, "학습 범위: WebSocket 연결 상태, 채팅 메시지 목록, 입력창과 전송 흐름")

    add_heading(document, "11.2 예제 구성과 패키지 확인", 2)
    add_body(document, "이번 예제는 Expo 기본 환경 위에서 동작하며, 별도의 외부 채팅 라이브러리 없이 React Native 기본 컴포넌트와 브라우저 표준 WebSocket API를 그대로 사용한다. 덕분에 연결 상태, 수신 이벤트, 전송 로직을 코드 수준에서 직접 살펴보기 좋다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "채팅 UI는 처음부터 복잡한 라이브러리를 붙이기보다, 기본 컴포넌트로 한 번 직접 만들어 보는 편이 상태 흐름을 이해하는 데 훨씬 도움이 된다. 이후 라이브러리를 도입하더라도 어떤 부분을 맡기는지 판단하기 쉬워진다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "11.3 예제 프로젝트 실행", 1)
    add_body(document, "예제를 실행하면 상단에는 연결 상태 배지가, 중앙에는 메시지 목록이, 하단에는 입력창과 전송 버튼이 나타난다. 아직 서버가 연결되지 않았더라도 전송 버튼을 눌러 로컬 메시지가 목록에 쌓이는 흐름을 확인할 수 있고, 서버를 준비한 뒤에는 같은 구조를 그대로 재사용할 수 있다.")
    add_command_block(
        document,
        [
            "cd chapters/11_WebSocket으로_실시간_채팅_화면_만들기/examples/expo-chat-ui",
            "npx expo start --android",
        ],
    )
    add_body(document, "실행 후 `WS_URL` 값을 자신의 로컬 서버 주소에 맞게 바꾸면 된다. 예를 들어 같은 와이파이에 연결된 개발 PC에서 8080 포트를 열었다면, 예제의 주소를 그 PC의 로컬 IP로 맞춰 주면 된다. 12장에서 실제 서버를 만들고 나면 이 주소를 사용해 양방향 통신이 완성된다.")

    add_heading(document, "11.4 WebSocket 이벤트를 화면 상태에 반영하기", 2)
    add_body(document, "이 장의 핵심은 WebSocket 이벤트를 React 상태와 연결하는 것이다. `onopen`에서는 연결 성공 상태를 기록하고, `onmessage`에서는 수신 메시지를 배열에 추가하며, 전송 버튼에서는 입력창 텍스트를 로컬 목록에 먼저 반영한 뒤 소켓으로 전송한다. 이 구조를 이해하면 채팅뿐 아니라 알림 스트림, 실시간 대시보드 같은 기능에도 그대로 응용할 수 있다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "WebSocket 서버 주소를 상수로 분리해 두면, 개발 서버 주소를 바꾸거나 12장에서 실제 서버를 붙일 때 수정 지점이 명확해진다."),
            ("②", "`useRef`에 소켓 객체를 저장하면 렌더링이 다시 일어나더라도 같은 연결 객체를 유지하면서 전송 버튼에서 안전하게 접근할 수 있다."),
            ("③", "`useEffect` 안에서 `new WebSocket()`을 호출하면 화면이 처음 열릴 때 연결을 시작하는 흐름이 만들어진다. 이 시점에 `onopen`, `onmessage`, `onerror`, `onclose` 이벤트도 함께 연결한다."),
            ("④", "서버에서 메시지가 도착하면 `event.data`를 읽어 새 메시지 객체로 바꾼 뒤 상태 배열 끝에 붙인다. 실시간 기능은 결국 외부 이벤트를 화면 상태로 변환하는 과정이라고 이해하면 쉽다."),
            ("⑤", "전송 버튼을 누르면 서버 응답을 기다리기 전에 내 메시지를 먼저 화면에 추가한다. 이렇게 해야 사용자는 버튼이 즉시 반응했다고 느끼고, 채팅 UI도 자연스럽게 보인다."),
            ("⑥", "소켓 상태가 실제로 열려 있을 때만 `send()`를 호출한다. 연결이 아직 준비되지 않았다면 로컬 안내 메시지를 추가해 현재 상태를 사용자에게 설명할 수 있다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명을 함께 읽어 보면, 이번 장은 단순히 채팅 말풍선 모양을 만드는 연습이 아니라 네이티브 앱에서 실시간 이벤트를 어떻게 다뤄야 하는지 배우는 과정이라는 점이 더 분명해진다.")

    document.add_page_break()

    add_heading(document, "11.5 결과 화면", 1)
    add_body(document, "아래 이미지는 연결 상태 배지, 시스템 안내 메시지, 상대 메시지, 내 메시지, 입력창과 전송 버튼이 한 화면에 정리된 예시다. 12장에서 서버를 붙이면 이 레이아웃을 그대로 유지한 채 메시지가 실제로 오가게 된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 11-1. WebSocket 채팅 화면 예시", width_cm=7.6)

    add_heading(document, "11.6 정리", 2)
    add_body(document, "이번 장에서는 WebSocket 서버가 아직 없어도 채팅 UI와 상태 흐름을 충분히 먼저 설계할 수 있다는 점을 확인했다. 연결 상태를 텍스트와 배지로 표현하고, 메시지 목록을 상태 배열로 관리하며, 입력창과 전송 버튼을 묶어 실시간 채팅 화면의 뼈대를 완성했다.")
    add_body(document, "다음 장에서는 Node.js 기반의 간단한 WebSocket 서버를 만들고, 이번 장의 클라이언트를 실제 서버와 연결해 여러 사용자가 메시지를 주고받는 구조로 확장할 것이다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 세 가지다. 앱이 열릴 때 연결 상태가 바뀌는지, 전송 버튼을 누르면 내 메시지가 즉시 말풍선으로 추가되는지, 서버 연결 전후에 화면 설명 문구를 어떻게 바꿔 줄 수 있는지 살펴보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
