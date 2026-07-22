import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "12_채팅_서버_구축과_메시지_브로드캐스트"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "websocket-chat-server"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "12_채팅_서버_구축과_메시지_브로드캐스트.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "12장 채팅 서버 구축과 메시지 브로드캐스트", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "Node.js WebSocket 서버로 여러 사용자의 메시지를 한 번에 전달하는 구조 익히기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "Node.js와 ws 패키지로 가장 단순한 WebSocket 채팅 서버를 만들고, 한 클라이언트의 메시지를 연결된 모든 사용자에게 다시 보내는 브로드캐스트 흐름을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "이번 장은 인증이나 데이터베이스 없이 서버의 핵심 역할인 접속 관리와 메시지 전달에 집중한다. 11장의 React Native 클라이언트를 실제 서버와 연결하는 단계까지 함께 정리한다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch12_websocket_server.jpg"
    annotated_server_js = "\n".join(
        [
            "const { WebSocketServer, WebSocket } = require('ws'); ①",
            "const PORT = 8080;",
            "const wss = new WebSocketServer({ port: PORT }); ②",
            "",
            "wss.on('connection', (socket) => { ③",
            "  const userId = `user-${Date.now()}`;",
            "  sendSystemMessage(`${userId} 님이 입장했습니다.`);",
            "",
            "  socket.on('message', (rawMessage) => { ④",
            "    const text = rawMessage.toString().trim();",
            "    broadcastMessage(`${userId}: ${text}`);",
            "  });",
            "",
            "  socket.on('close', () => {",
            "    sendSystemMessage(`${userId} 님이 퇴장했습니다.`);",
            "  });",
            "});",
            "",
            "function broadcastMessage(message) { ⑤",
            "  wss.clients.forEach((client) => {",
            "    if (client.readyState === WebSocket.OPEN) {",
            "      client.send(message); ⑥",
            "    }",
            "  });",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "12장 채팅 서버 구축과 메시지 브로드캐스트")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "12.1 채팅 서버는 왜 필요한가", 1)
    add_body(document, "11장에서는 React Native 화면 안에서 메시지 목록과 연결 상태를 다뤘지만, 여러 사용자가 실제로 서로의 메시지를 보려면 중간에서 메시지를 받아 다시 전달해 주는 서버가 필요하다. 채팅 서버는 누가 접속했는지 기억하고, 한 사용자가 보낸 텍스트를 나머지 사용자에게 브로드캐스트하는 역할을 맡는다.")
    add_body(document, "이번 장은 인증, 저장소, 방 분리 같은 복잡한 기능을 잠시 내려놓고, 가장 핵심적인 세 가지 흐름만 다룬다. 첫째는 접속 이벤트, 둘째는 메시지 수신 이벤트, 셋째는 전체 사용자에게 다시 보내는 브로드캐스트다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Node.js 서버 + 11장의 Android 채팅 앱")
    add_bullet(document, "학습 범위: WebSocket 서버 생성, 클라이언트 접속 관리, 전체 메시지 전파")

    add_heading(document, "12.2 서버 예제와 패키지 준비", 2)
    add_body(document, "Node.js에서 WebSocket 서버를 가장 빠르게 시작하는 방법 중 하나가 ws 패키지다. 브라우저 WebSocket과 개념은 비슷하지만 서버 쪽에서 연결 목록을 직접 다룰 수 있어서, 브로드캐스트 구조를 설명하기에 적합하다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "처음부터 데이터베이스나 로그인까지 넣으면 채팅 서버의 핵심 흐름이 흐려진다. 우선은 접속과 전달만 되는 최소 서버를 만든 뒤, 이후 장이나 별도 프로젝트에서 인증과 저장 기능을 덧붙이는 편이 학습 효율이 좋다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "12.3 서버 실행과 11장 앱 연결", 1)
    add_body(document, "서버 예제는 Node.js 콘솔에서 실행하고, 11장 앱의 `WS_URL`을 같은 네트워크의 PC IP로 맞춰 연결하면 된다. 예를 들어 서버 PC 주소가 `192.168.0.10`이고 포트가 `8080`이라면, 앱에는 `ws://192.168.0.10:8080`을 입력하면 된다.")
    add_command_block(
        document,
        [
            "cd chapters/12_채팅_서버_구축과_메시지_브로드캐스트/examples/websocket-chat-server",
            "npm install",
            "npm start",
        ],
    )
    add_body(document, "서버를 실행한 뒤 11장 앱을 두 대 이상 열어 메시지를 보내 보면, 한 기기에서 입력한 내용이 다른 기기에도 그대로 나타나는지 확인할 수 있다. 콘솔에는 누가 접속했고 어떤 메시지가 들어왔는지 로그가 쌓인다.")

    add_heading(document, "12.4 접속 목록을 이용해 메시지 브로드캐스트하기", 2)
    add_body(document, "이 장의 핵심은 서버가 `wss.clients` 목록을 이용해 현재 열려 있는 모든 소켓에 같은 메시지를 보내는 부분이다. 특정 사용자에게만 답장을 보내는 것이 아니라, 접속 중인 전체 사용자에게 같은 내용을 다시 전송하기 때문에 채팅방의 기본 구조가 만들어진다.")
    add_code_block(document, "server.js", annotated_server_js)
    add_code_notes(
        document,
        [
            ("①", "`ws` 패키지에서 서버 생성자와 연결 상태 상수를 함께 가져온다. 브로드캐스트를 할 때는 `WebSocket.OPEN` 값을 비교하는 부분이 중요하다."),
            ("②", "포트 번호를 정하고 `WebSocketServer`를 만들면 이 서버가 클라이언트 연결을 기다리기 시작한다. 11장의 앱은 바로 이 주소로 접속한다."),
            ("③", "새 사용자가 들어오면 `connection` 이벤트가 실행된다. 여기서 사용자 식별자와 입장 안내 메시지를 만들면, 누가 채팅방에 들어왔는지 전체 사용자에게 보여 줄 수 있다."),
            ("④", "클라이언트에서 보낸 원본 데이터는 버퍼일 수 있으므로 문자열로 바꾸고 공백을 정리한 뒤 사용한다. 이후 브로드캐스트 메시지 형식으로 감싸서 전체 사용자에게 다시 전달한다."),
            ("⑤", "`broadcastMessage()`는 이 장의 중심 함수다. 서버가 직접 모든 연결을 순회하면서 같은 메시지를 보내는 구조를 한 곳에 모아 두면 코드가 훨씬 읽기 쉬워진다."),
            ("⑥", "소켓이 아직 열린 상태일 때만 `send()`를 호출해야 오류를 줄일 수 있다. 그래서 브로드캐스트에는 항상 연결 상태 확인이 따라붙는다고 이해하면 좋다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명까지 함께 읽으면, 서버는 단순히 메시지를 저장하는 장소가 아니라 연결된 사용자 목록을 관리하고 실시간 이벤트를 다시 배포하는 중심 허브라는 점이 분명해진다.")

    document.add_page_break()

    add_heading(document, "12.5 실행 결과", 1)
    add_body(document, "아래 예시는 왼쪽 콘솔에서 접속과 메시지 로그가 쌓이고, 오른쪽 두 클라이언트 화면에서 같은 대화가 동시에 보이는 모습을 정리한 것이다. 11장 앱과 이번 장 서버가 합쳐지면 비로소 실시간 채팅 기능이 완성된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 12-1. WebSocket 서버와 두 클라이언트의 브로드캐스트 예시", width_cm=12.5)

    add_heading(document, "12.6 정리", 2)
    add_body(document, "이번 장에서는 Node.js와 ws 패키지로 가장 단순한 WebSocket 채팅 서버를 만들고, 접속 이벤트와 메시지 수신 이벤트를 이용해 전체 사용자에게 다시 메시지를 보내는 구조를 완성했다.")
    add_body(document, "이제 React Native 클라이언트와 서버가 모두 준비되었으므로, 다음 장부터는 푸시 메시지나 지도 같은 다른 네이티브 기능을 다룰 때도 '클라이언트 상태 + 서버 이벤트'라는 관점으로 구조를 더 쉽게 이해할 수 있다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 서버가 실행되면 콘솔에 포트가 표시되는지, 앱 두 대가 모두 연결되는지, 한 기기 메시지가 다른 기기에도 보이는지, 기기를 닫았을 때 퇴장 로그가 남는지 확인해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
