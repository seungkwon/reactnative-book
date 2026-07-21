import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "02_기본_컴포넌트로_화면_구성하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-basic-components"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "02_기본_컴포넌트로_화면_구성하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "2장. 기본 컴포넌트로 화면 구성하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "Image, TextInput, Pressable로 프로필 카드 화면 만들기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "View, Text, Image, TextInput, Pressable을 조합해 작은 프로필 카드 화면을 만들고, 컴포넌트와 스타일의 역할을 구분해 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 화면 구성 자체에 집중한다. 버튼 클릭이나 입력 상태 변경 같은 상호작용 로직은 다음 장에서 다룬다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch2_profile_card.jpg"
    annotated_app_js = "\n".join(
        [
            "import { StatusBar } from 'expo-status-bar';",
            "import { Image, Pressable, SafeAreaView, StyleSheet, Text, TextInput, View } from 'react-native';",
            "",
            "const avatarImage = require('./assets/profile-card-avatar.png');",
            "",
            "export default function App() {  ①",
            "  return (",
            "    <SafeAreaView style={styles.screen}>  ②",
            "      <View style={styles.card}>",
            "        <Text style={styles.badge}>Chapter 02</Text>",
            "        <Image  ③",
            "          source={avatarImage}",
            "          style={styles.avatar}",
            "        />",
            "        <Text style={styles.name}>코딩하는 리액트 개발자</Text>",
            "        <Text style={styles.role}>React Native UI Starter</Text>",
            "        <TextInput  ④",
            "          style={styles.input}",
            "          placeholder=\"별명을 입력해 보세요\"",
            "          placeholderTextColor=\"#94a3b8\"",
            "        />",
            "        <Pressable style={styles.button}>  ⑤",
            "          <Text style={styles.buttonText}>프로필 완성하기</Text>",
            "        </Pressable>",
            "      </View>",
            "      <StatusBar style=\"dark\" />",
            "    </SafeAreaView>",
            "  );",
            "}",
            "",
            "const styles = StyleSheet.create({",
            "  screen: {  ⑥",
            "    flex: 1,",
            "    backgroundColor: '#e0f2fe',",
            "    justifyContent: 'center',",
            "    paddingHorizontal: 24,",
            "  },",
            "  card: {",
            "    backgroundColor: '#ffffff',",
            "    borderRadius: 28,",
            "    paddingHorizontal: 24,",
            "    paddingVertical: 28,",
            "    alignItems: 'center',",
            "  },",
            "  avatar: {",
            "    width: 108,",
            "    height: 108,",
            "    borderRadius: 54,",
            "  },",
            "});",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "2장. 기본 컴포넌트로 화면 구성하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "2.1 왜 기본 컴포넌트를 먼저 익혀야 할까", 1)
    add_body(document, "React Native 화면은 여러 기본 컴포넌트를 조합해 만든다. 첫 장에서 View와 Text만으로 첫 화면을 확인했다면, 이제는 이미지와 입력창, 버튼까지 더해 실제 앱 화면에 가까운 구성을 경험할 차례다.")
    add_body(document, "이번 장의 핵심은 기능을 많이 넣는 것이 아니라, 화면을 이루는 요소마다 어떤 역할이 있는지 또렷하게 구분하는 것이다. 이를 통해 독자는 앞으로 더 복잡한 화면을 만들 때도 구조를 먼저 설계하는 습관을 익힐 수 있다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 화면 구성과 스타일 분리")

    add_heading(document, "2.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 별도의 외부 UI 라이브러리를 추가하지 않고, React Native가 기본 제공하는 컴포넌트만으로 카드형 화면을 구성했다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "처음부터 외부 UI 라이브러리에 의존하기보다 기본 컴포넌트만으로 화면을 구성해 보면, 각 요소의 책임과 스타일 동작 방식을 훨씬 명확하게 이해할 수 있다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "2.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 02 예제는 chapter 01과 같은 Expo 실행 흐름을 유지한다. 독자는 같은 방식으로 앱을 열되, 이번에는 한 단계 더 완성도 있는 UI가 화면에 나타나는 것을 확인하면 된다.")
    add_command_block(
        document,
        [
            "cd chapters/02_기본_컴포넌트로_화면_구성하기/examples/expo-basic-components",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "프로젝트를 실행하면 연한 하늘색 배경 위에 흰색 카드가 놓인 프로필 화면이 보인다. 이 화면 안에서 독자는 이미지, 텍스트, 입력창, 버튼이 각각 어떤 방식으로 배치되는지 관찰할 수 있다.")

    add_heading(document, "2.4 프로필 카드 화면 구현", 2)
    add_body(document, "이번 예제의 중심은 하나의 카드 안에 여러 기본 컴포넌트를 질서 있게 배치하는 것이다. 복잡한 비즈니스 로직 없이도 화면 구성이 얼마나 풍부해질 수 있는지 보여 주기에 적합한 예제다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "App 컴포넌트는 chapter 02 화면 전체를 반환하는 진입점이다. 지금은 상태 관리가 없기 때문에 구조를 읽는 연습에 집중하기 좋다."),
            ("②", "SafeAreaView는 기기의 안전 영역 안에서 콘텐츠가 보이도록 돕는다. 상단 상태 표시줄이나 노치 영역과 겹치지 않게 화면을 시작하는 기본 습관으로 이해하면 좋다."),
            ("③", "Image는 텍스트만 있던 화면에 시각적 중심점을 만든다. 프로필 카드처럼 이미지가 중요한 화면에서는 레이아웃의 인상을 결정하는 핵심 요소가 된다."),
            ("④", "TextInput은 사용자의 입력이 들어올 자리를 미리 마련해 두는 컴포넌트다. 아직 값을 처리하지 않더라도, 입력 UI가 어떤 속성으로 구성되는지 익히는 데 큰 의미가 있다."),
            ("⑤", "Pressable은 버튼처럼 눌릴 수 있는 영역을 만들 때 자주 쓰는 컴포넌트다. 다음 장에서 상호작용을 연결하기 전, 먼저 모양과 배치부터 익히기에 적합하다."),
            ("⑥", "screen 스타일은 배경색, 중앙 정렬, 좌우 여백처럼 화면 전체의 분위기를 정하는 공통 설정이다. 상위 컨테이너 스타일을 먼저 이해하면 하위 카드 레이아웃도 훨씬 쉽게 읽힌다."),
        ],
    )
    add_body(document, "이 예제에서 가장 중요한 포인트는 컴포넌트가 늘어날수록 JSX 구조와 스타일 객체를 더 의식적으로 분리해야 한다는 점이다. 화면의 요소 수가 조금만 늘어나도, 정돈된 구조와 이름이 유지보수 난이도를 크게 바꾼다.")

    document.add_page_break()

    add_heading(document, "2.5 화면 구성에서 눈여겨볼 점", 1)
    add_body(document, "카드형 화면은 모바일 앱에서 매우 자주 등장하는 패턴이다. 이번 장에서는 카드 내부 요소를 세로 방향으로 차례대로 쌓고, 텍스트 계층과 버튼 위치를 분리해 사용자의 시선 흐름이 자연스럽게 이어지도록 구성했다.")
    add_body(document, "또한 입력창과 버튼을 카드 하단에 배치해, 단순한 소개 화면을 넘어 앞으로 상호작용을 붙일 수 있는 형태의 기반 UI로 확장했다. 즉 chapter 02는 단순히 예쁜 화면을 만드는 것이 아니라, 다음 장으로 이어지는 발판을 준비하는 역할도 한다.")

    if ui_screen.exists():
        add_image(document, ui_screen, "그림 2-1. 기본 컴포넌트로 구성한 프로필 카드 화면", width_cm=7.6)
    else:
        add_note_box(
            document,
            "캡처 예정",
            "현재 원고에는 실행 화면 캡처가 아직 포함되지 않았다. Android 에뮬레이터 실행 후 프로필 카드 화면을 캡처해 artifacts 폴더에 추가하면 그림 영역을 자연스럽게 보강할 수 있다.",
            "FFFBEB",
            "B45309",
        )

    add_heading(document, "2.6 정리", 2)
    add_body(document, "이 장에서는 React Native 기본 컴포넌트를 조합해 카드형 화면을 만들었다. 독자는 View와 Text를 넘어 Image, TextInput, Pressable이 실제 화면 안에서 어떤 역할을 맡는지 직접 확인할 수 있다.")
    add_body(document, "다음 장에서는 여기에 상태를 연결해 입력값을 화면에 반영하거나 버튼을 눌렀을 때 텍스트가 바뀌는 방식까지 확장하면 학습 흐름이 자연스럽게 이어진다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 컴포넌트가 늘어나도 JSX 구조와 StyleSheet 역할이 분리된 상태로 유지되는지, 그리고 각 요소가 카드 화면 안에서 의도한 위치에 배치되는지다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
