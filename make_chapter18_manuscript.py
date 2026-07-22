import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "18_WebView로_하이브리드_화면_연동하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "rn-webview-hybrid"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "18_WebView로_하이브리드_화면_연동하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "18장 WebView로 하이브리드 화면 연동하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "웹 화면과 네이티브 화면이 메시지를 주고받는 하이브리드 앱의 기본 구조 익히기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "react-native-webview로 웹 화면을 임베드하고, WebView와 네이티브 앱이 postMessage와 injectJavaScript로 양방향 메시지를 주고받는 구조를 만든다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "이번 장은 기본 임베드와 양방향 통신에 집중한다. 로그인 쿠키, 파일 업로드, 외부 결제, 다중 창 제어 같은 고급 주제는 이후 확장 영역으로 남겨 둔다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch18_webview_hybrid.jpg"
    annotated_app_js = "\n".join(
        [
            "import { WebView } from 'react-native-webview'; ①",
            "",
            "const htmlSource = `... window.ReactNativeWebView.postMessage(JSON.stringify({ type: 'WEB_ACTION' })) ...`; ②",
            "",
            "const handleMessage = (event) => { ③",
            "  const payload = JSON.parse(event.nativeEvent.data);",
            "  setLastMessage(payload.title);",
            "};",
            "",
            "const sendToWeb = () => {",
            "  webviewRef.current?.injectJavaScript(`window.showNativeMessage('React Native에서 보낸 확인 메시지'); true;`); ④",
            "};",
            "",
            "<WebView",
            "  ref={webviewRef}",
            "  originWhitelist={['*']} ⑤",
            "  source={{ html: htmlSource }} ⑥",
            "  onMessage={handleMessage}",
            "/>",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "18장 WebView로 하이브리드 화면 연동하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "18.1 왜 하이브리드 화면이 여전히 중요한가", 1)
    add_body(document, "모든 화면을 완전한 네이티브로만 만드는 것이 항상 정답은 아니다. 이미 구축된 웹 관리자, 약관 화면, 이벤트 페이지, 실험 기능을 빠르게 앱 안으로 들여와야 할 때 WebView 기반 하이브리드 구조가 여전히 유효하다.")
    add_body(document, "2026년 7월 22일 기준 react-native-webview 공식 저장소는 최신 릴리스를 2026년 2월 27일의 `v13.16.1`로 표시하고 있으며, 공식 가이드는 `source`, `onMessage`, `injectedJavaScript`를 대표적인 기본 흐름으로 설명한다. 따라서 이번 장은 가장 많이 쓰이는 하이브리드 패턴인 '웹이 네이티브에 알리고, 네이티브가 웹에 다시 명령을 보내는 구조'를 중심으로 정리한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: React Native CLI Android 또는 iOS 앱")
    add_bullet(document, "학습 범위: WebView 임베드, postMessage, injectJavaScript, 최근 메시지 상태 표시")

    add_heading(document, "18.2 예제 패키지와 기본 구성", 2)
    add_body(document, "WebView는 더 이상 React Native 코어에 포함되지 않기 때문에 공식 커뮤니티 패키지를 따로 설치해야 한다. 공식 Getting Started 문서도 `npm install --save react-native-webview` 또는 `yarn add react-native-webview`를 안내하고 있다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "공식 Reference 문서는 인라인 HTML을 `source={{ html: ... }}`로 넣을 때 `originWhitelist={['*']}` 설정이 필요할 수 있다고 안내한다. 책 예제도 이 흐름을 그대로 사용한다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "18.3 웹과 네이티브가 서로 메시지를 주고받게 만들기", 1)
    add_body(document, "하이브리드 앱에서 가장 중요한 것은 단순 렌더링이 아니라 통신 지점이다. 웹 화면이 버튼 클릭 사실을 네이티브에 알리고, 네이티브는 확인 메시지를 다시 웹 화면에 주입해 상태를 바꿔 준다. 이 양방향 흐름이 있으면 결제 완료, 인증 상태 갱신, 브릿지 이벤트 같은 실제 하이브리드 기능으로 확장하기 쉬워진다.")
    add_command_block(
        document,
        [
            "cd chapters/18_WebView로_하이브리드_화면_연동하기/examples/rn-webview-hybrid",
            "npm install",
            "npm run android",
        ],
    )
    add_body(document, "예제를 실행하면 상단에는 최근 수신 메시지 카드가, 하단에는 WebView 안의 웹 화면이 보인다. 웹 버튼을 누르면 네이티브 카드가 갱신되고, 네이티브 버튼을 누르면 WebView 내부 문구가 바뀌는 흐름을 확인할 수 있다.")

    add_heading(document, "18.4 WebView 통신 코드를 읽는 방법", 2)
    add_body(document, "이번 장의 핵심은 렌더링보다 통신 구조다. 웹은 `window.ReactNativeWebView.postMessage()`로 메시지를 보내고, 네이티브는 `onMessage`로 받아 상태를 바꾸며, 다시 `injectJavaScript()`로 웹의 함수를 호출한다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "react-native-webview 공식 패키지에서 WebView 컴포넌트를 가져온다. 하이브리드 화면의 진입점은 항상 이 컴포넌트다."),
            ("②", "웹 안에서는 `window.ReactNativeWebView.postMessage()`가 네이티브로 통하는 기본 브리지다. 공식 가이드의 대표 패턴도 이 흐름을 중심으로 설명한다."),
            ("③", "네이티브 쪽에서는 `onMessage` 이벤트로 웹의 메시지를 받는다. 여기서 JSON을 해석해 React 상태로 바꾸면 하이브리드 이벤트가 앱 UI에 반영된다."),
            ("④", "반대로 네이티브에서 웹으로 무언가 지시하고 싶다면 `injectJavaScript()`를 사용한다. 이 메서드는 웹 내부 함수를 직접 호출하거나 DOM 상태를 바꿀 때 자주 쓰인다."),
            ("⑤", "인라인 HTML을 안전하게 로드하려면 허용 출처 설정을 확인해야 한다. 공식 문서도 `originWhitelist`와 `source` 조합을 함께 설명한다."),
            ("⑥", "`source`에는 외부 URL뿐 아니라 인라인 HTML도 넣을 수 있다. 초반 학습 단계에서는 HTML 문자열을 직접 넣는 편이 통신 구조를 이해하기 쉽다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명까지 함께 읽으면, 이번 장은 단순히 웹을 앱 안에 띄우는 장이 아니라 웹과 네이티브의 경계를 설계하는 장이라는 점이 선명해진다.")

    document.add_page_break()

    add_heading(document, "18.5 결과 화면", 1)
    add_body(document, "아래 예시는 상단의 최근 수신 메시지 카드, 네이티브 전송 버튼, 그리고 WebView 안쪽 웹 패널이 한 화면에 배치된 모습이다. 이 장까지 마치면 네이티브 기능과 웹 화면 연동을 함께 다루는 교재 흐름이 완성된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 18-1. WebView 하이브리드 메시지 통신 화면", width_cm=7.6)

    add_heading(document, "18.6 정리", 2)
    add_body(document, "이번 장에서는 react-native-webview로 웹 화면을 임베드하고, 웹에서 네이티브로 메시지를 보내고 네이티브에서 웹으로 다시 메시지를 주입하는 기본 하이브리드 구조를 만들었다. 이 흐름은 약관 동의, 이벤트 페이지, 레거시 관리자 연동 같은 실제 업무 시나리오로 바로 이어질 수 있다.")
    add_body(document, "이 장으로 책의 확장 네이티브 기능 파트도 한 사이클을 마무리했다. 이제 독자는 카메라, 채팅, 푸시, 지도, 소셜 로그인, WebView 하이브리드까지 모바일 앱 실무에서 자주 만나는 연동 주제를 한 흐름으로 경험한 셈이다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 웹 버튼을 누르면 네이티브 카드가 바뀌는지, 네이티브 버튼을 누르면 웹 문구가 바뀌는지, 인라인 HTML이 정상적으로 보이는지, 메시지 JSON을 상태로 잘 바꾸는지 확인해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
