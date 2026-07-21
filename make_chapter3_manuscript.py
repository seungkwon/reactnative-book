import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "03_useState로_화면_바꾸기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-state-profile"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "03_useState로_화면_바꾸기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "3장. useState로 화면 바꾸기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "입력값과 버튼 상태를 화면에 즉시 반영하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "useState를 이용해 입력값과 버튼 상태를 관리하고, 상태가 바뀌면 화면 문구와 버튼 모양이 어떻게 함께 갱신되는지 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 useState 하나만 집중적으로 다룬다. 네트워크 요청, useEffect, 서버 연동 같은 주제는 포함하지 않는다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch3_state_profile.jpg"
    annotated_app_js = "\n".join(
        [
            "import { StatusBar } from 'expo-status-bar';",
            "import { useState } from 'react';  ①",
            "import { Image, Pressable, SafeAreaView, StyleSheet, Text, TextInput, View } from 'react-native';",
            "",
            "const avatarImage = require('./assets/profile-card-avatar.png');",
            "",
            "export default function App() {",
            "  const [nickname, setNickname] = useState('코딩하는 리액트 개발자');  ②",
            "  const [isFollowing, setIsFollowing] = useState(false);",
            "",
            "  const buttonLabel = isFollowing ? '팔로잉 취소하기' : '팔로우 시작하기';",
            "  const statusMessage = isFollowing  ③",
            "    ? `${nickname} 님의 새 글 알림을 받고 있습니다.`",
            "    : `${nickname} 님을 팔로우하면 새 소식을 빠르게 확인할 수 있습니다.`;",
            "",
            "  return (",
            "    <SafeAreaView style={styles.screen}>",
            "      <View style={styles.card}>",
            "        <Text style={styles.name}>{nickname}</Text>  ④",
            "        <Text style={styles.description}>{statusMessage}</Text>",
            "        <TextInput",
            "          style={styles.input}",
            "          value={nickname}",
            "          onChangeText={setNickname}  ⑤",
            "        />",
            "        <Pressable",
            "          style={[styles.button, isFollowing && styles.buttonActive]}",
            "          onPress={() => setIsFollowing((current) => !current)}  ⑥",
            "        >",
            "          <Text style={styles.buttonText}>{buttonLabel}</Text>",
            "        </Pressable>",
            "      </View>",
            "    </SafeAreaView>",
            "  );",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "3장. useState로 화면 바꾸기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "3.1 왜 상태가 필요한가", 1)
    add_body(document, "chapter 02에서는 여러 컴포넌트를 조합해 화면을 구성하는 데 집중했다. 하지만 실제 앱은 화면이 한 번 그려지고 끝나지 않는다. 입력값이 바뀌거나 버튼을 누를 때마다 내용이 다시 달라져야 한다.")
    add_body(document, "React의 상태는 바로 이 변화를 다루는 핵심 도구다. 이번 장에서는 가장 기본적인 훅인 useState만으로도 화면이 얼마나 즉각적으로 달라질 수 있는지 확인한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 입력 상태와 토글 상태")

    add_heading(document, "3.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 새로운 라이브러리는 추가하지 않고 React가 기본 제공하는 useState만으로 상호작용을 구현한다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "useState 예제는 가능한 한 작고 눈에 보이게 만드는 편이 좋다. 입력값과 버튼 상태처럼 즉시 확인 가능한 요소를 쓰면 상태 개념을 훨씬 빠르게 체감할 수 있다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "3.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 03 예제는 chapter 02의 프로필 카드 화면을 기반으로 한다. 이번에는 이름 입력과 팔로우 버튼이 상태와 연결되어, 사용자의 행동에 따라 화면이 실시간으로 달라진다.")
    add_command_block(
        document,
        [
            "cd chapters/03_useState로_화면_바꾸기/examples/expo-state-profile",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 입력창에 다른 이름을 입력해 보고, 버튼을 눌러 상태 메시지와 버튼 색이 바뀌는지 확인해 보자. 이 짧은 상호작용 안에 useState의 핵심 흐름이 모두 담겨 있다.")

    add_heading(document, "3.4 useState로 상태 연결하기", 2)
    add_body(document, "이번 예제는 두 개의 상태를 사용한다. 하나는 사용자가 입력하는 별명이고, 다른 하나는 팔로우 여부다. 상태가 바뀌면 React Native는 해당 화면을 다시 렌더링하므로, 개발자는 변경된 값을 JSX에 연결해 주기만 하면 된다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "useState를 가져오면 함수형 컴포넌트 안에서 변경 가능한 값을 상태로 관리할 수 있다. React Native 실습에서 가장 먼저 익혀야 할 기본 훅이다."),
            ("②", "nickname과 isFollowing은 각각 입력값과 팔로우 여부를 저장하는 상태다. 화면이 단순해 보여도, 실제로는 두 값이 바뀔 때마다 다시 그려질 준비를 하고 있는 셈이다."),
            ("③", "상태값을 바탕으로 문구를 계산하면 JSX 안에서 조건문을 길게 쓰지 않아도 된다. 화면에 보일 텍스트를 미리 정리해 두는 습관은 코드 가독성을 높이는 데 도움이 된다."),
            ("④", "nickname 상태를 Text 컴포넌트에 연결해 두었기 때문에, 입력창에서 값을 바꾸는 순간 화면 제목도 함께 바뀐다. 이 즉시 반영이 상태 관리의 핵심 경험이다."),
            ("⑤", "TextInput의 value와 onChangeText를 상태와 연결하면 입력 필드가 완전한 제어 컴포넌트가 된다. 입력값을 코드가 알고 있기 때문에 이후 검증이나 저장 기능도 쉽게 확장할 수 있다."),
            ("⑥", "Pressable의 onPress에서 이전 값을 뒤집도록 작성하면 토글 동작을 간단하게 만들 수 있다. 버튼을 누를 때마다 상태, 문구, 스타일이 한 번에 바뀌는 흐름을 꼭 눈으로 확인해 보는 것이 좋다."),
        ],
    )
    add_body(document, "useState를 배우는 가장 좋은 방법은 복잡한 이론보다 작은 변화를 여러 번 확인하는 것이다. 입력값이 제목에 반영되고, 버튼을 누를 때 문구가 바뀌는 경험만으로도 상태의 의미가 꽤 선명해진다.")

    document.add_page_break()

    add_heading(document, "3.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 별명을 `리액트 쌤`으로 바꾸고, 팔로우 버튼을 눌러 활성 상태가 된 결과다. chapter 02와 비교하면 화면 구성은 비슷하지만, 이번에는 사용자의 입력과 클릭이 화면 내용에 직접 영향을 준다는 점이 가장 큰 차이점이다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 3-1. useState로 입력값과 버튼 상태를 반영한 화면", width_cm=7.6)

    add_heading(document, "3.6 정리", 2)
    add_body(document, "이 장에서는 useState를 이용해 입력값과 버튼 상태를 화면에 연결했다. 독자는 정적인 UI에서 한 단계 나아가, 사용자 행동에 반응하는 화면이 어떻게 만들어지는지 직접 확인할 수 있다.")
    add_body(document, "다음 장에서는 리스트 데이터나 여러 개의 아이템을 렌더링하는 방식으로 확장하면, 상태 관리와 화면 구조가 함께 커지는 흐름을 자연스럽게 이어갈 수 있다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 입력창의 값이 제목에 즉시 반영되는지, 버튼을 누를 때 문구와 버튼 색이 함께 바뀌는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
