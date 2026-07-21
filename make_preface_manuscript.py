from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "00_서문과_학습_로드맵"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "00_서문과_학습_로드맵.docx"


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_roadmap_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["장", "핵심 주제", "배우는 내용"]
    cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        set_cell_shading(cells[idx], "DCE6F1")
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, color="1A365D", size=10.5)

    rows = [
        ("1장", "Hello World", "Expo 프로젝트 생성과 첫 실행 흐름 익히기"),
        ("2장", "기본 컴포넌트", "View, Text, Image, TextInput, Pressable 조합"),
        ("3장", "useState", "입력값과 버튼 상태를 화면 변화에 연결"),
        ("4장", "FlatList", "여러 데이터를 목록으로 렌더링"),
        ("5장", "목록 추가", "입력값으로 새 항목을 배열에 추가"),
        ("6장", "검색", "검색어에 따라 목록 결과 실시간 필터링"),
        ("7장", "정렬", "정렬 기준에 따라 목록 순서 변경"),
        ("8장", "상세 보기", "선택한 항목의 상세 정보 연결"),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            add_run(cells[idx].paragraphs[0], value, size=10.5)


def main() -> None:
    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "서문과 학습 로드맵")

    for section in document.sections:
        configure_page(section)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "서문과 학습 로드맵", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "React Native 첫 실습부터 목록과 상세 화면까지 이어지는 8단계 흐름", color="555555", size=13)

    add_note_box(
        document,
        "이 책의 방향",
        "이 책은 React Native를 처음 접하는 독자가 작은 화면 예제를 차근차근 확장하면서 모바일 UI와 상태 관리의 기본 흐름을 익히도록 구성했다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "읽는 방법",
        "장별 예제를 직접 실행하고, 코드 스니펫 아래 해설을 함께 읽는 방식으로 진행하면 학습 효과가 가장 좋다. 중간 장으로 건너뛰기보다 1장부터 순서대로 따라가는 편을 권장한다.",
        "FFF7ED",
        "C2410C",
    )

    add_heading(document, "0.1 이 책이 다루는 범위", 1)
    add_body(document, "이 책은 React Native 입문자가 모바일 화면을 직접 만들고, 상태를 연결하고, 목록을 다루고, 선택한 데이터를 상세하게 보여 주는 흐름까지 경험하도록 돕는다. 복잡한 빌드 환경이나 네이티브 심화 주제보다, 화면을 보면서 바로 이해할 수 있는 실습 중심 흐름에 초점을 맞췄다.")
    add_body(document, "특히 1장부터 8장까지의 예제는 서로 단절되지 않고 자연스럽게 이어진다. 첫 Hello World에서 시작해 컴포넌트, 상태, 목록, 검색, 정렬, 상세 보기로 확장되도록 설계했기 때문에, 독자는 매 장마다 조금씩 더 넓은 화면 구조를 경험하게 된다.")

    add_heading(document, "0.2 어떤 독자를 위한 책인가", 2)
    add_bullet(document, "React Native를 처음 배우는 입문자")
    add_bullet(document, "웹 프론트엔드 경험은 있지만 모바일 UI 흐름은 익숙하지 않은 개발자")
    add_bullet(document, "짧은 예제로 상태와 목록 처리 감각을 빠르게 익히고 싶은 독자")

    add_heading(document, "0.3 실습 환경과 진행 방식", 1)
    add_body(document, "이 책의 예제는 Windows 환경과 Android 실행 기준으로 정리했다. Expo를 사용해 프로젝트를 빠르게 시작하고, 각 장의 예제를 실제 화면으로 확인하는 데 집중한다.")
    add_body(document, "모든 장에는 예제 앱, notes 문서, 결과 화면 캡처, 워드 원고가 함께 준비되어 있다. 따라서 독자는 코드만 보지 않고 결과 화면과 설명을 동시에 비교하면서 학습할 수 있다.")

    add_heading(document, "0.4 장별 학습 로드맵", 1)
    add_body(document, "아래 표는 책 전체 흐름을 한눈에 보여 준다. 앞 장에서 배운 개념이 다음 장의 기반이 되므로, 각 장의 위치와 역할을 함께 보면 전체 학습 구조를 더 쉽게 이해할 수 있다.")
    add_roadmap_table(document)

    add_heading(document, "0.5 공부 팁", 2)
    add_body(document, "React Native를 처음 배울 때는 한 번에 많은 기능을 넣기보다, 화면이 바뀌는 이유를 짧은 코드로 자주 확인하는 편이 훨씬 효과적이다. 이 책의 예제도 바로 그 점을 염두에 두고 구성했다.")
    add_bullet(document, "코드를 읽기 전에 먼저 결과 화면을 보고, 어떤 변화가 일어나는지 예상해 본다.")
    add_bullet(document, "코드 스니펫의 원문자 표시와 해설을 함께 읽으며 각 부분의 역할을 구분한다.")
    add_bullet(document, "장별 예제를 직접 실행해 보고, 한두 군데 값을 바꿔 화면 변화를 확인해 본다.")
    add_bullet(document, "검색, 정렬, 상세 보기처럼 이전 장을 확장하는 흐름을 특히 주의 깊게 본다.")

    add_heading(document, "0.6 마치며", 1)
    add_body(document, "이 책의 목표는 React Native의 모든 기능을 한 권에 담는 것이 아니라, 입문자가 실제로 손에 잡히는 화면 예제를 통해 기본기를 빠르게 세우도록 돕는 것이다. 한 장씩 차근차근 따라오다 보면, 단순한 텍스트 화면에서 시작해 상호작용이 있는 목록과 상세 구조까지 자연스럽게 도달할 수 있다.")
    add_body(document, "이제 1장부터 차례대로 실습을 시작해 보자. 첫 화면을 직접 띄우는 순간부터 책 전체의 흐름이 훨씬 또렷하게 잡히기 시작할 것이다.")

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
