import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "14_구글_맵으로_현재_위치와_마커_표시하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "rn-current-location-map"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "14_구글_맵으로_현재_위치와_마커_표시하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "14장 구글 맵으로 현재 위치와 마커 표시하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "위치 권한과 좌표를 지도 중심, 사용자 위치, 마커에 함께 반영하는 지도 기본 구조 익히기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "react-native-maps로 지도 화면을 만들고 현재 위치 좌표를 가져와 지도 중심과 마커, 상태 카드에 반영한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "이번 장은 현재 위치 표시와 단일 마커에 집중한다. 주소 검색과 좌표 변환, 경로 탐색, 네이버 맵 연동은 다음 장 이후에 확장한다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch14_current_location_map.jpg"
    annotated_app_js = "\n".join(
        [
            "import MapView, { Marker, PROVIDER_GOOGLE } from 'react-native-maps'; ①",
            "import Geolocation from 'react-native-geolocation-service'; ②",
            "",
            "const defaultRegion = { latitude: 37.5665, longitude: 126.978, latitudeDelta: 0.01, longitudeDelta: 0.01 };",
            "",
            "const moveToCurrentLocation = async () => {",
            "  const granted = await requestLocationPermission(); ③",
            "  Geolocation.getCurrentPosition((position) => { ④",
            "    const nextRegion = {",
            "      latitude: position.coords.latitude,",
            "      longitude: position.coords.longitude,",
            "      latitudeDelta: 0.01,",
            "      longitudeDelta: 0.01,",
            "    };",
            "    setCurrentRegion(nextRegion);",
            "    setMarkerPosition(nextRegion); ⑤",
            "    mapRef.current?.animateToRegion(nextRegion, 600); ⑥",
            "  });",
            "};",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "14장 구글 맵으로 현재 위치와 마커 표시하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "14.1 지도 기능의 기본 흐름", 1)
    add_body(document, "지도 기능은 처음 보면 복잡해 보이지만 실제 흐름은 단순하다. 먼저 지도를 화면에 띄우고, 위치 권한을 요청하고, 좌표를 받아와서 지도 중심이나 마커 위치로 반영하면 된다. 이번 장은 이 세 단계를 한 화면에서 연결하는 데 집중한다.")
    add_body(document, "react-native-maps 공식 설치 문서는 Android에서 Google Maps API 키가 필요하다고 안내하고 있으며, MapView 문서는 `showsUserLocation` 사용 전에 런타임 위치 권한이 필요하다고 설명한다. 따라서 지도 기능은 화면 코드만이 아니라 네이티브 설정과 권한 흐름까지 같이 이해해야 한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: React Native CLI Android 앱")
    add_bullet(document, "학습 범위: 위치 권한, 현재 좌표 획득, 지도 중심 이동, 마커 표시")

    add_heading(document, "14.2 패키지와 네이티브 준비", 2)
    add_body(document, "이번 예제는 지도를 렌더링하는 react-native-maps와 현재 위치를 가져오는 react-native-geolocation-service를 함께 사용한다. 지도는 화면을 담당하고, 위치 라이브러리는 좌표를 얻는 역할을 맡는다.")
    add_dependency_table(document, package_json)
    add_note_box(
        document,
        "TIP",
        "Android에서는 `AndroidManifest.xml`에 위치 권한을 추가하고 Google Maps API 키를 메타데이터로 넣어야 한다. 공식 설치 문서 기준으로 `<meta-data android:name=\"com.google.android.geo.API_KEY\" ... />` 항목이 필요하다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "14.3 현재 위치 버튼과 마커 연결", 1)
    add_body(document, "버튼을 누르면 먼저 권한을 요청하고, 허용되면 현재 좌표를 얻은 뒤, 그 좌표를 상태에 저장해 지도 중심과 마커를 동시에 업데이트한다. 이렇게 해야 사용자는 현재 위치가 숫자로도 보이고, 지도에서도 바로 확인할 수 있다.")
    add_command_block(
        document,
        [
            "cd chapters/14_구글_맵으로_현재_위치와_마커_표시하기/examples/rn-current-location-map",
            "npm install",
            "npm run android",
        ],
    )
    add_body(document, "실행 전에는 Android Google Maps API 키와 위치 권한 설정이 끝나 있어야 한다. 키가 없으면 지도 타일이 비어 보일 수 있고, 권한이 없으면 현재 위치 버튼을 눌러도 좌표를 가져오지 못한다.")

    add_heading(document, "14.4 위치 좌표를 지도 상태로 반영하기", 2)
    add_body(document, "이 장의 핵심은 위치 좌표를 단순히 받아오는 데서 끝내지 않고, 그것을 지도 UI 상태로 자연스럽게 연결하는 부분이다. 권한 요청, 좌표 획득, 상태 반영, 카메라 이동이 하나의 흐름으로 이어진다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "react-native-maps 공식 문서 기준으로 Android에서는 Google Maps를 사용하므로 `PROVIDER_GOOGLE` 지정이 일반적이다. Marker 컴포넌트로 지도 위 특정 지점을 표시할 수 있다."),
            ("②", "지도 컴포넌트만으로는 현재 좌표를 가져올 수 없기 때문에 위치 전용 라이브러리를 함께 사용한다. 여기서는 geolocation-service가 좌표 획득을 맡는다."),
            ("③", "위치 좌표를 요청하기 전에 먼저 권한을 확인해야 한다. 위치 권한은 네이티브 기능의 첫 관문이므로 버튼 흐름 시작점에 두는 편이 이해하기 쉽다."),
            ("④", "좌표를 실제로 받는 지점이다. 성공 콜백 안에서 위도와 경도를 읽어 다음 지역 정보를 만들면 이후 상태 반영을 한 번에 처리할 수 있다."),
            ("⑤", "받아온 좌표를 마커 상태에 넣으면 지도 위 표시점이 바로 생긴다. 좌표 값과 시각적 표시를 같은 데이터로 연결하는 전형적인 지도 패턴이다."),
            ("⑥", "animateToRegion을 호출하면 지도 중심이 현재 위치로 부드럽게 이동한다. 단순히 숫자만 바뀌는 것이 아니라 사용자가 이동 결과를 시각적으로 체감할 수 있게 된다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명까지 함께 보면, 이번 장은 지도 컴포넌트를 띄우는 수준을 넘어 위치라는 네이티브 데이터를 지도 상태와 UI로 연결하는 첫 단계라는 점이 분명해진다.")

    document.add_page_break()

    add_heading(document, "14.5 결과 화면", 1)
    add_body(document, "아래 예시는 권한 상태 카드, 현재 좌표, 현재 위치 버튼, 그리고 지도 위 마커가 한 화면에 정리된 모습이다. 다음 장에서는 여기에 주소 검색과 좌표 변환 흐름을 이어 붙일 예정이다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 14-1. 현재 위치와 마커를 표시한 지도 화면", width_cm=7.6)

    add_heading(document, "14.6 정리", 2)
    add_body(document, "이번 장에서는 react-native-maps와 위치 라이브러리를 이용해 현재 위치를 지도 중심과 마커에 반영하는 기본 구조를 만들었다. 지도 기능은 결국 권한 요청, 좌표 획득, UI 반영의 세 단계로 이해할 수 있다는 점을 확인했다.")
    add_body(document, "다음 장에서는 여기서 한 걸음 더 나아가 주소 검색 결과를 좌표로 바꾸고, 사용자가 입력한 위치를 지도 위에 표시하는 흐름으로 확장한다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 현재 위치 버튼을 누르면 권한 상태가 바뀌는지, 좌표 숫자가 갱신되는지, 지도 중심이 이동하는지, 마커가 같은 위치에 표시되는지 확인해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
