import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "01_엑스포로_Hello_World_만들기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-hello-world"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "01_엑스포로_Hello_World_만들기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "1장. 엑스포로 Hello World 만들기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "Windows 환경에서 Android 에뮬레이터로 실행하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "Expo 프로젝트를 만들고 Hello World 화면을 작성한 뒤, Windows에서 Android 에뮬레이터로 실행하여 실제 화면을 캡처한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 Windows 환경과 Android 실행만 다룬다. 웹 빌드와 iOS 관련 절차는 포함하지 않는다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    bundling_screen = ARTIFACTS_DIR / "ch1_screen.png"
    final_screen = ARTIFACTS_DIR / "ch1_hello_world_final.png"
    annotated_app_js = "\n".join(
        [
            "import { StatusBar } from 'expo-status-bar';",
            "import { StyleSheet, Text, View } from 'react-native';",
            "",
            "export default function App() {  ①",
            "  return (",
            "    <View style={styles.container}>  ②",
            "      <Text style={styles.badge}>Chapter 01</Text>",
            "      <Text style={styles.title}>Hello World</Text>  ③",
            "      <Text style={styles.subtitle}>엑스포로 만드는 첫 리액트 네이티브 앱</Text>",
            "      <Text style={styles.description}>  ④",
            "        Windows 환경에서 Expo 프로젝트를 만들고 첫 화면을 확인합니다.",
            "      </Text>",
            "      <StatusBar style=\"auto\" />  ⑤",
            "    </View>",
            "  );",
            "}",
            "",
            "const styles = StyleSheet.create({",
            "  container: {  ⑥",
            "    flex: 1,",
            "    backgroundColor: '#f3f7fb',",
            "    alignItems: 'center',",
            "    justifyContent: 'center',",
            "    paddingHorizontal: 24,",
            "  },",
            "  badge: {",
            "    fontSize: 14,",
            "    fontWeight: '700',",
            "    color: '#2563eb',",
            "    marginBottom: 12,",
            "    textTransform: 'uppercase',",
            "    letterSpacing: 1,",
            "  },",
            "  title: {",
            "    fontSize: 36,",
            "    fontWeight: '800',",
            "    color: '#0f172a',",
            "  },",
            "  subtitle: {",
            "    fontSize: 18,",
            "    fontWeight: '600',",
            "    color: '#1d4ed8',",
            "    marginTop: 10,",
            "    textAlign: 'center',",
            "  },",
            "  description: {",
            "    marginTop: 16,",
            "    fontSize: 15,",
            "    lineHeight: 22,",
            "    color: '#334155',",
            "    textAlign: 'center',",
            "  },",
            "});",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "1장. 엑스포로 Hello World 만들기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "1.1 Expo와 이번 장의 목표", 1)
    add_body(document, "Expo는 React Native 프로젝트를 빠르게 시작할 수 있게 도와주는 도구 모음이다. 첫 장에서는 복잡한 네이티브 설정을 깊게 파고들기보다, 직접 화면을 만들고 실행해 보는 경험을 우선 확보하는 것이 중요하다.")
    add_body(document, "이번 실습은 Windows 환경을 기준으로 진행했고, Android 에뮬레이터에서 Hello World 화면을 실제로 실행해 확인했다. 독자는 프로젝트 생성, 코드 수정, 에뮬레이터 실행, 화면 확인이라는 기본 흐름을 한 번에 익힐 수 있다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터 Pixel_7")
    add_bullet(document, "실행 방식: Expo Go를 통한 Android 실행")

    add_heading(document, "1.2 준비 환경", 2)
    add_body(document, "실습에는 Node.js, npm, Java, Android SDK가 필요하다. 이 장의 목적은 Android 에뮬레이터에서 첫 화면을 확인하는 것이므로, 개발 서버가 실행되고 adb가 에뮬레이터를 인식하는지 확인하는 것이 핵심이다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "Windows 환경에서는 Android Studio와 SDK가 준비되어 있으면 Expo 앱을 비교적 쉽게 확인할 수 있다. 실제 휴대폰이 없어도 에뮬레이터만으로 학습 흐름을 이어갈 수 있다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "1.3 프로젝트 생성과 실행 명령", 1)
    add_body(document, "프로젝트는 blank 템플릿으로 생성했다. blank 템플릿은 불필요한 기본 예제가 거의 없어서 책의 첫 장에서 핵심 개념만 설명하기 좋다.")
    add_command_block(
        document,
        [
            "npx create-expo-app@latest expo-hello-world --template blank --yes",
            "cd expo-hello-world",
            "npx expo start --android",
        ],
    )
    add_body(document, "npx expo start --android 명령을 실행하면 Metro 번들러가 올라오고, 연결된 Android 에뮬레이터 또는 실제 폰에서 Expo Go로 프로젝트가 열린다. 이 장에서는 Pixel_7 에뮬레이터를 사용해 결과를 확인했다.")

    add_heading(document, "1.4 Hello World 화면 구현", 2)
    add_body(document, "기본으로 생성된 App.js는 매우 단순한 구조다. 여기에 장 번호, 제목, 부제, 설명 문장을 추가해 첫 장에 어울리는 화면으로 바꾸었다. 독자는 이 과정을 통해 View와 Text, StyleSheet의 기본 사용법을 자연스럽게 익힐 수 있다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "App 컴포넌트는 이 예제의 진입점이다. React Native는 이 함수가 반환하는 JSX를 기준으로 첫 화면을 그리므로, 화면 구성을 읽을 때 가장 먼저 확인해야 하는 위치다."),
            ("②", "최상위 View는 화면 전체를 감싸는 컨테이너다. 자식 요소들의 정렬 기준이 여기서 결정되기 때문에, 레이아웃이 흐트러질 때는 이 블록의 스타일을 먼저 확인하는 습관이 중요하다."),
            ("③", "Hello World 제목은 사용자가 앱을 실행했을 때 가장 먼저 인식하는 핵심 메시지다. 큰 글자와 굵은 두께를 적용해 화면의 시선을 모으고, 예제가 정상적으로 렌더링되었는지 바로 판단할 수 있게 한다."),
            ("④", "설명 문장은 단순한 장식이 아니라 화면의 목적을 보완하는 역할을 한다. 실습 환경과 학습 목표를 함께 보여 주기 때문에, 독자는 지금 만든 앱이 무엇을 확인하려는 예제인지 즉시 이해할 수 있다."),
            ("⑤", "StatusBar는 기기 상단 상태 표시줄의 표시 스타일을 현재 화면과 자연스럽게 맞춘다. 예제 규모가 작더라도 운영체제 UI와 앱 화면이 어색하게 분리되지 않도록 챙기는 기본 설정이라고 이해하면 좋다."),
            ("⑥", "container 스타일은 중앙 정렬, 배경색, 좌우 여백처럼 화면의 전체 분위기를 결정하는 속성을 묶어 둔 곳이다. 레이아웃과 시각 스타일을 컴포넌트 바깥으로 분리해 두면 화면 구조와 디자인 역할이 나뉘어 유지보수가 쉬워진다."),
        ],
    )
    add_body(document, "핵심은 화면을 이루는 요소를 작게 나누고, 스타일을 별도 객체로 관리하는 습관을 처음부터 익히는 것이다. 예제가 단순할수록 각 요소의 역할이 더 선명하게 드러난다.")

    document.add_page_break()

    add_heading(document, "1.5 에뮬레이터 실행 과정", 1)
    add_body(document, "실행 초기에는 Expo Go가 프로젝트를 번들링하는 화면이 나타난다. 이 장면은 개발 서버와 에뮬레이터가 정상적으로 연결되었는지 확인하는 중요한 중간 지점이다.")
    if bundling_screen.exists():
        add_image(document, bundling_screen, "그림 1-1. Android 에뮬레이터에서 Expo Go가 프로젝트를 번들링하는 화면", width_cm=7.6)
    add_note_box(
        document,
        "확인 포인트",
        "에뮬레이터 상단에 Bundling 메시지가 보이면 Metro와 에뮬레이터 연결이 정상적으로 진행되고 있다는 뜻이다.",
        "FFFBEB",
        "B45309",
    )

    add_heading(document, "1.6 최종 실행 화면", 2)
    add_body(document, "번들링이 끝나면 Hello World 화면이 실제로 표시된다. 아래 캡처는 Windows 환경에서 Pixel_7 Android 에뮬레이터로 실행한 최종 결과 화면이다.")
    if final_screen.exists():
        add_image(document, final_screen, "그림 1-2. Android 에뮬레이터에서 실행된 Hello World 최종 화면", width_cm=7.6)
    add_body(document, "이 단계에서 독자는 코드 수정 결과가 모바일 화면에 어떻게 반영되는지 직접 확인하게 된다. 첫 장에서는 이 경험 자체가 가장 중요하며, 이후의 장은 이 기본 실행 흐름 위에 기능을 쌓아 가는 방식으로 이어진다.")

    document.add_page_break()

    add_heading(document, "1.7 정리", 1)
    add_body(document, "이 장에서는 Expo를 이용해 React Native Hello World 프로젝트를 만들고, Windows 환경에서 Android 에뮬레이터로 실제 실행 결과를 확인했다. 단순한 예제이지만 개발 환경을 준비하고 실행 흐름을 검증하는 데 필요한 핵심 단계가 모두 담겨 있다.")
    add_body(document, "다음 장에서는 화면 요소를 더 세분화하거나 입력 요소를 추가하면서 컴포넌트 구조를 조금씩 확장해 볼 수 있다. 지금 단계에서는 프로젝트 생성, App.js 수정, Android 실행, 화면 캡처 네 가지를 확실하게 익히는 것이 가장 중요하다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 npx expo start --android 실행 후 에뮬레이터에서 Hello World 화면이 열리는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
