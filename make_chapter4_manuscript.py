import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "04_FlatList로_목록_화면_만들기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-flatlist-profiles"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "04_FlatList로_목록_화면_만들기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "4장. FlatList로 목록 화면 만들기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "여러 개의 프로필 카드를 반복 렌더링하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "FlatList를 이용해 여러 개의 프로필 데이터를 목록으로 출력하고, renderItem과 keyExtractor의 역할을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 로컬 배열 데이터를 기반으로 한 목록 렌더링에 집중한다. 서버에서 데이터를 불러오는 과정은 포함하지 않는다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch4_flatlist_profiles.jpg"
    annotated_app_js = "\n".join(
        [
            "import { StatusBar } from 'expo-status-bar';",
            "import { useState } from 'react';",
            "import { FlatList, Image, Pressable, SafeAreaView, StyleSheet, Text, View } from 'react-native';  ①",
            "",
            "const initialProfiles = [  ②",
            "  { id: '1', name: '리액트 쌤', role: 'React Native State Starter', following: true },",
            "  { id: '2', name: '모바일 메이트', role: 'UI Component Explorer', following: false },",
            "  { id: '3', name: '앱 빌더', role: 'FlatList Practice Partner', following: false },",
            "];",
            "",
            "export default function App() {",
            "  const [profiles, setProfiles] = useState(initialProfiles);  ③",
            "",
            "  const toggleFollow = (targetId) => {",
            "    setProfiles((currentProfiles) =>",
            "      currentProfiles.map((profile) =>",
            "        profile.id === targetId ? { ...profile, following: !profile.following } : profile",
            "      )",
            "    );",
            "  };",
            "",
            "  const renderItem = ({ item }) => (  ④",
            "    <View style={styles.card}>",
            "      <Text style={styles.name}>{item.name}</Text>",
            "      <Pressable onPress={() => toggleFollow(item.id)}>  ⑤",
            "        <Text>{item.following ? '팔로잉' : '팔로우'}</Text>",
            "      </Pressable>",
            "    </View>",
            "  );",
            "",
            "  return (",
            "    <SafeAreaView style={styles.screen}>",
            "      <FlatList  ⑥",
            "        data={profiles}",
            "        keyExtractor={(item) => item.id}",
            "        renderItem={renderItem}",
            "      />",
            "    </SafeAreaView>",
            "  );",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "4장. FlatList로 목록 화면 만들기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "4.1 왜 목록 화면이 중요한가", 1)
    add_body(document, "실제 앱은 단일 카드 하나만 보여 주는 경우보다 여러 데이터를 나열하는 경우가 훨씬 많다. 상품 목록, 채팅 목록, 뉴스 목록, 알림 목록처럼 대부분의 모바일 화면은 결국 리스트를 다루게 된다.")
    add_body(document, "React Native에서는 이런 반복 렌더링을 효율적으로 처리하기 위해 FlatList를 자주 사용한다. 이번 장에서는 chapter 03에서 다루었던 상태 개념을 유지한 채, 같은 구조를 여러 항목으로 확장해 본다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 배열 데이터와 FlatList 렌더링")

    add_heading(document, "4.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 새로운 외부 라이브러리 없이 React Native 기본 제공 컴포넌트와 FlatList만으로 추천 프로필 목록 화면을 완성한다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "목록 화면을 처음 배울 때는 API 호출부터 시작하기보다, 먼저 로컬 배열 데이터로 renderItem과 keyExtractor 구조를 충분히 익히는 편이 훨씬 안정적이다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "4.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 04 예제는 추천 프로필 세 개를 세로 목록으로 보여 준다. 각 항목마다 팔로우 버튼이 있고, 버튼을 누르면 그 항목의 상태만 개별적으로 바뀐다.")
    add_command_block(
        document,
        [
            "cd chapters/04_FlatList로_목록_화면_만들기/examples/expo-flatlist-profiles",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "실행 후 스크롤 가능한 목록이 보이면 FlatList가 정상 동작하는 것이다. 각 버튼을 눌러 항목마다 라벨과 색이 바뀌는지 확인해 보자.")

    add_heading(document, "4.4 FlatList 구조 이해하기", 2)
    add_body(document, "이번 예제는 배열 데이터, 상태, renderItem, FlatList 네 요소가 함께 맞물린다. 데이터는 화면에 표시할 원본이고, renderItem은 각 항목을 카드 모양으로 변환하며, FlatList는 이를 실제 스크롤 목록으로 그린다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "FlatList를 import하면 여러 항목을 효율적으로 렌더링할 수 있다. 단순 map 렌더링보다 모바일 목록 화면에 더 적합한 기본 도구라고 이해하면 된다."),
            ("②", "initialProfiles 배열은 화면에 표시할 원본 데이터다. 지금은 세 개뿐이지만, 같은 구조의 데이터가 많아져도 같은 방식으로 확장할 수 있다."),
            ("③", "profiles 상태는 목록 전체를 관리한다. 항목 하나의 팔로우 상태가 바뀌더라도 결국은 배열 안의 한 원소가 변경되는 것이므로, 목록 상태 전체를 다루는 감각이 중요하다."),
            ("④", "renderItem은 배열의 각 항목을 받아 실제 카드 UI로 바꿔 주는 함수다. 목록 UI를 설계할 때 가장 자주 손보게 되는 핵심 블록이다."),
            ("⑤", "항목별 버튼은 item.id를 기준으로 어떤 프로필을 바꿀지 결정한다. 목록 화면에서는 전체 상태를 건드리기보다, 정확히 어느 항목이 바뀌는지 식별하는 과정이 특히 중요하다."),
            ("⑥", "FlatList는 data, keyExtractor, renderItem 세 속성만 알아도 기본 구조를 이해할 수 있다. 이 세 가지가 연결되면 배열 데이터가 스크롤 가능한 실제 화면으로 바뀐다."),
        ],
    )
    add_body(document, "목록 화면에서 가장 자주 실수하는 부분은 키 값 관리와 항목별 상태 변경이다. 이번 예제처럼 작은 구조에서부터 식별자와 상태 갱신 흐름을 정확히 잡아 두면 이후의 복잡한 목록에서도 훨씬 안정적으로 작업할 수 있다.")

    document.add_page_break()

    add_heading(document, "4.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 추천 프로필 목록을 FlatList로 출력한 결과다. 첫 번째 항목은 이미 팔로잉 상태이고, 나머지 항목은 팔로우 버튼을 눌러 상태를 바꿀 수 있도록 구성했다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 4-1. FlatList로 렌더링한 추천 프로필 목록 화면", width_cm=7.6)

    add_heading(document, "4.6 정리", 2)
    add_body(document, "이 장에서는 FlatList를 이용해 여러 프로필 카드를 목록으로 렌더링했다. 독자는 단일 화면 중심의 예제에서 벗어나, 실제 앱에서 매우 자주 등장하는 목록 UI 구조를 직접 다뤄 볼 수 있다.")
    add_body(document, "다음 장에서는 이 목록 데이터를 입력과 연결해 새 항목을 추가하거나, 필터링하는 방식으로 확장하면 상태 관리와 리스트 렌더링이 한층 더 입체적으로 이어진다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 FlatList가 세 개의 프로필 카드를 안정적으로 렌더링하는지, 그리고 항목별 버튼이 각각 독립적으로 상태를 바꾸는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
