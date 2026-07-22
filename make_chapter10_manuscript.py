import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "10_선택한_이미지를_프로필_화면에_적용하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-profile-image"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "10_선택한_이미지를_프로필_화면에_적용하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "10장. 선택한 이미지를 프로필 화면에 적용하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "이미지 선택 결과를 실제 프로필 아바타 UI에 연결하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "갤러리나 카메라에서 받은 이미지 URI를 실제 프로필 카드의 아바타와 설명 문구에 연결해, 네이티브 기능 결과가 화면 UI 일부로 반영되는 흐름을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 이미지 선택 결과를 UI에 적용하는 데 집중한다. 서버 업로드나 프로필 저장 API 연동은 이후 단계에서 확장할 수 있도록 남겨 둔다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch10_profile_image.jpg"
    annotated_app_js = "\n".join(
        [
            "import * as ImagePicker from 'expo-image-picker';",
            "const defaultAvatar = require('./assets/profile-card-avatar.png');  ①",
            "",
            "export default function App() {",
            "  const [selectedImage, setSelectedImage] = useState(null);  ②",
            "  const [sourceLabel, setSourceLabel] = useState('기본 프로필 이미지가 적용되어 있습니다.');",
            "",
            "  const pickFromGallery = async () => {",
            "    const permission = await ImagePicker.requestMediaLibraryPermissionsAsync();",
            "    const result = await ImagePicker.launchImageLibraryAsync({ allowsEditing: true });  ③",
            "",
            "    if (!result.canceled && result.assets[0]) {",
            "      setSelectedImage(result.assets[0].uri);  ④",
            "      setSourceLabel('갤러리에서 이미지를 선택했습니다.');",
            "    }",
            "  };",
            "",
            "  return (",
            "    <Image",
            "      source={selectedImage ? { uri: selectedImage } : defaultAvatar}  ⑤",
            "      style={styles.avatar}",
            "    />",
            "  );",
            "",
            "  <Text style={styles.helperText}>{sourceLabel}</Text>  ⑥",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "10장. 선택한 이미지를 프로필 화면에 적용하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "10.1 왜 이미지 선택 결과를 화면에 적용해야 할까", 1)
    add_body(document, "chapter 09에서는 카메라와 갤러리에서 이미지를 가져오는 기본 흐름을 익혔다. 하지만 실제 앱에서는 이미지를 단순히 받아 오는 데서 끝나지 않고, 그 결과를 프로필 사진이나 게시물 썸네일처럼 사용자에게 보이는 실제 UI에 반영해야 한다.")
    add_body(document, "이번 장에서는 선택된 이미지 URI를 프로필 카드의 아바타에 연결한다. 이를 통해 네이티브 기능 결과가 앱 화면의 한 요소로 자연스럽게 편입되는 흐름을 경험할 수 있다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터 또는 실제 Android 기기")
    add_bullet(document, "학습 범위: 이미지 선택 결과를 프로필 UI에 적용")

    add_heading(document, "10.2 준비 환경과 패키지", 2)
    add_body(document, "이번 장도 expo-image-picker를 그대로 사용한다. 다른 점은 미리보기 박스 대신 실제 프로필 카드 구조를 만들고, 선택된 이미지가 그 카드의 아바타로 적용되도록 연결한다는 점이다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "네이티브 기능을 배울 때는 결과를 단순 로그나 임시 미리보기에만 두지 말고, 실제 앱 UI 구조에 붙여 보는 편이 훨씬 실전적인 감각을 길러 준다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "10.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 10 예제는 프로필 카드 한 장을 화면 중앙에 배치하고, 갤러리 또는 카메라 선택 결과를 이 카드의 프로필 이미지로 반영한다. 기본 이미지를 보여 주다가 사용자가 선택한 이미지를 새 아바타로 바꾸는 흐름을 확인하는 것이 목표다.")
    add_command_block(
        document,
        [
            "cd chapters/10_선택한_이미지를_프로필_화면에_적용하기/examples/expo-profile-image",
            "npx expo install expo-image-picker",
            "npx expo start --android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 갤러리에서 이미지를 고르거나 카메라로 촬영해 보자. 이미지가 성공적으로 선택되면 프로필 카드의 둥근 아바타 이미지가 즉시 새 사진으로 바뀌게 된다.")

    add_heading(document, "10.4 이미지 URI를 프로필 카드에 연결하기", 2)
    add_body(document, "이번 예제의 핵심은 선택 결과를 단순한 미리보기가 아니라 실제 프로필 UI에 연결하는 것이다. 즉 URI를 어디에 저장할지보다, 저장된 URI를 어떤 시점에 어떤 컴포넌트의 source로 연결할지가 더 중요하다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "defaultAvatar는 사용자가 아직 이미지를 선택하지 않았을 때 보여 줄 기본 아바타다. 초기 상태를 분명히 만들어 두면 선택 전후 변화가 훨씬 선명하게 보인다."),
            ("②", "selectedImage 상태는 선택된 사용자 이미지를 저장한다. 이 값 하나가 바뀌는 것만으로도 프로필 카드의 핵심 시각 요소가 달라질 수 있다."),
            ("③", "갤러리 선택 로직은 9장과 거의 같지만, 이번에는 결과를 미리보기 박스가 아니라 프로필 카드 구조와 연결한다는 점이 다르다."),
            ("④", "URI를 상태에 넣는 순간부터 React Native는 이 값을 실제 이미지 소스로 사용할 수 있다. 네이티브 결과와 UI 상태의 연결이 가장 분명하게 드러나는 지점이다."),
            ("⑤", "source 속성에서 selectedImage가 있으면 사용자 이미지를, 없으면 기본 아바타를 사용하도록 분기한다. 이 패턴은 프로필 사진, 썸네일, 첨부 이미지 등 다양한 화면에서 반복적으로 쓰인다."),
            ("⑥", "도움말 문구도 함께 바꾸면 사용자는 현재 화면이 어떤 상태인지 더 쉽게 이해할 수 있다. 기능 동작만큼이나 상태 설명을 함께 보여 주는 습관이 UI 완성도를 높인다."),
        ],
    )
    add_body(document, "이 장을 통해 독자는 네이티브 기능을 다루는 목적이 단순한 기능 호출이 아니라, 그 결과를 사용자가 인식할 수 있는 실제 화면 경험으로 연결하는 데 있다는 점을 확인할 수 있다.")

    document.add_page_break()

    add_heading(document, "10.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 사용자가 이미지를 선택한 뒤 프로필 카드의 아바타가 바뀐 예시를 표현한 것이다. 기본 아바타 대신 선택 결과가 프로필 정체성을 보여 주는 핵심 요소로 작동한다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 10-1. 선택한 이미지를 프로필 카드에 적용한 화면", width_cm=7.6)

    add_heading(document, "10.6 정리", 2)
    add_body(document, "이 장에서는 카메라 또는 갤러리에서 가져온 이미지를 프로필 화면에 직접 적용했다. 독자는 네이티브 기능 결과를 실제 앱 UI에 연결하는 기본 패턴을 익히게 된다.")
    add_body(document, "다음 장부터는 화면 단위 예제를 넘어 네트워크와 실시간 통신이 필요한 채팅 기능으로 확장해, React Native 앱이 서버와 연결되는 흐름까지 다뤄 볼 수 있다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 이미지를 선택한 뒤 프로필 카드의 둥근 아바타가 실제로 바뀌는지, 그리고 상태 문구도 함께 갱신되는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
