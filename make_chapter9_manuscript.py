import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "09_카메라와_이미지_갤러리_연동하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-image-picker-demo"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "09_카메라와_이미지_갤러리_연동하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "9장. 카메라와 이미지 갤러리 연동하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "촬영과 갤러리 선택 결과를 React Native 화면에 반영하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "expo-image-picker를 이용해 카메라 촬영과 갤러리 선택 결과를 받아 오고, 선택된 이미지 URI를 화면 미리보기와 연결하는 방법을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 이미지 선택 자체와 미리보기 반영에 집중한다. 업로드 서버 연동이나 영구 저장은 다음 장에서 프로필 화면과 연결하며 확장한다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch9_image_picker.jpg"
    annotated_app_js = "\n".join(
        [
            "import * as ImagePicker from 'expo-image-picker';  ①",
            "import { useState } from 'react';",
            "",
            "export default function App() {",
            "  const [selectedImage, setSelectedImage] = useState(null);  ②",
            "  const [sourceLabel, setSourceLabel] = useState('아직 이미지를 선택하지 않았습니다.');",
            "",
            "  const pickFromGallery = async () => {",
            "    const permission = await ImagePicker.requestMediaLibraryPermissionsAsync();  ③",
            "    if (!permission.granted) {",
            "      setSourceLabel('갤러리 접근 권한이 필요합니다.');",
            "      return;",
            "    }",
            "",
            "    const result = await ImagePicker.launchImageLibraryAsync({",
            "      allowsEditing: true,",
            "    });  ④",
            "",
            "    if (!result.canceled && result.assets[0]) {",
            "      setSelectedImage(result.assets[0].uri);  ⑤",
            "      setSourceLabel('갤러리에서 이미지를 선택했습니다.');",
            "    }",
            "  };",
            "",
            "  const takePhoto = async () => {",
            "    const permission = await ImagePicker.requestCameraPermissionsAsync();",
            "    const result = await ImagePicker.launchCameraAsync({});  ⑥",
            "  };",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "9장. 카메라와 이미지 갤러리 연동하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "9.1 왜 카메라와 갤러리 연동이 중요한가", 1)
    add_body(document, "실제 모바일 앱에서는 텍스트 입력만큼이나 이미지 선택과 촬영이 자주 등장한다. 프로필 사진 변경, 리뷰 이미지 첨부, 문서 촬영, 게시물 업로드처럼 디바이스 카메라와 갤러리는 사용자 경험의 핵심 기능이 되는 경우가 많다.")
    add_body(document, "이번 장에서는 Expo 환경에서 카메라와 갤러리 연동을 가장 간단한 형태로 시작한다. 핵심은 사진 파일 자체보다, 권한 요청과 선택 결과를 받아 화면에 연결하는 기본 흐름을 익히는 것이다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터 또는 실제 Android 기기")
    add_bullet(document, "학습 범위: 권한 요청, 갤러리 선택, 카메라 촬영, 미리보기")

    add_heading(document, "9.2 준비 환경과 패키지", 2)
    add_body(document, "이번 장에서는 Expo 공식 패키지인 expo-image-picker를 사용한다. 이 패키지는 갤러리 접근과 카메라 실행을 비교적 단순한 API로 제공해, 네이티브 기능 입문 장에 적합하다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "카메라와 갤러리 기능을 처음 연동할 때는 업로드나 서버 저장까지 한 번에 붙이기보다, 먼저 권한 요청과 선택 결과 미리보기까지 확인하는 편이 구조를 훨씬 이해하기 쉽다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "9.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 09 예제는 두 개의 버튼을 제공한다. 하나는 갤러리에서 이미지를 선택하는 버튼이고, 다른 하나는 카메라를 열어 새 사진을 촬영하는 버튼이다. 두 경우 모두 결과 이미지 URI를 받아 미리보기 영역에 반영한다.")
    add_command_block(
        document,
        [
            "cd chapters/09_카메라와_이미지_갤러리_연동하기/examples/expo-image-picker-demo",
            "npx expo install expo-image-picker",
            "npx expo start --android",
        ],
    )
    add_body(document, "실행 후 먼저 갤러리 선택 버튼을 눌러 보고, 가능하다면 카메라 버튼도 눌러 촬영 흐름을 확인해 보자. 권한이 허용되면 결과 이미지가 곧바로 미리보기 카드에 반영된다.")

    add_heading(document, "9.4 권한 요청과 이미지 선택 흐름", 2)
    add_body(document, "이 예제의 핵심은 권한 요청, 선택 도구 실행, 결과 처리 세 단계를 차례로 연결하는 것이다. 네이티브 기능 연동은 대부분 이와 비슷한 패턴을 가지므로, 이번 장의 구조를 잘 이해해 두면 이후의 기능 연동에도 큰 도움이 된다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "expo-image-picker를 import하면 갤러리와 카메라를 여는 기능을 사용할 수 있다. Expo 기반 프로젝트에서는 네이티브 기능 입문용으로 가장 다루기 쉬운 선택지 중 하나다."),
            ("②", "selectedImage 상태는 사용자가 마지막으로 선택하거나 촬영한 이미지의 URI를 저장한다. 이 값을 화면과 연결하면 미리보기 영역이 즉시 바뀐다."),
            ("③", "갤러리 접근 전에는 반드시 권한을 확인해야 한다. 모바일 기능 연동에서 가장 먼저 체크해야 할 것은 권한 여부라는 점을 이번 장에서 분명히 익혀 두는 것이 좋다."),
            ("④", "launchImageLibraryAsync를 호출하면 시스템 갤러리가 열리고, 사용자가 선택한 결과가 비동기로 돌아온다. 이 단계에서 편집 옵션이나 비율 제한도 함께 지정할 수 있다."),
            ("⑤", "선택 결과에서 assets[0].uri를 꺼내 상태에 저장하면, 화면에서 바로 이미지로 사용할 수 있다. 네이티브 기능 결과를 UI 상태로 바꾸는 전형적인 패턴이다."),
            ("⑥", "카메라 흐름도 기본 구조는 동일하다. 권한을 받고, 카메라를 열고, 결과를 받아 상태에 넣는 3단계를 그대로 반복한다."),
        ],
    )
    add_body(document, "이번 장에서 중요한 것은 단순히 버튼이 카메라를 연다는 사실보다, 네이티브 기능 결과를 React Native 상태와 연결하는 감각이다. 이 흐름이 잡히면 파일 업로드, 프로필 사진 변경, 첨부 기능 같은 실전 예제로 자연스럽게 이어질 수 있다.")

    document.add_page_break()

    add_heading(document, "9.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 아직 이미지를 선택하지 않은 초기 상태의 예시다. 실제 기기 또는 에뮬레이터에서 버튼을 눌러 권한을 허용하고 이미지를 선택하면, 같은 위치에 선택한 이미지가 채워지게 된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 9-1. 카메라와 이미지 갤러리 연동 예제의 초기 화면", width_cm=7.6)

    add_heading(document, "9.6 정리", 2)
    add_body(document, "이 장에서는 expo-image-picker를 이용해 갤러리 선택과 카메라 촬영 결과를 React Native 화면에 연결했다. 독자는 네이티브 기능 연동의 기본 패턴인 권한 요청, 기능 실행, 결과 상태 반영 흐름을 직접 확인할 수 있다.")
    add_body(document, "다음 장에서는 여기서 얻은 이미지 URI를 프로필 화면과 연결해, 사용자가 선택한 사진이 실제 프로필 UI에 적용되도록 확장하면 흐름이 자연스럽게 이어진다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 갤러리나 카메라 결과가 선택된 뒤, 이미지 URI가 미리보기 화면에 즉시 반영되는지와 권한 거부 시 안내 문구가 바뀌는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
