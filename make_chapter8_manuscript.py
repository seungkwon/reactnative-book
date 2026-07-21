import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "08_목록에서_상세_화면으로_이동하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-profile-detail"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "08_목록에서_상세_화면으로_이동하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "8장. 목록에서 상세 화면으로 이동하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "선택한 항목의 자세한 정보를 별도 영역에 보여 주기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "목록 항목을 선택하면 선택 상태가 바뀌고, 그에 따라 상세 카드 내용이 함께 바뀌는 흐름을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 실제 네비게이션 라이브러리 대신 선택 상태 기반 상세 보기에 집중한다. 핵심은 화면 전환 개념보다 선택된 데이터와 상세 정보의 연결이다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch8_profile_detail.jpg"
    annotated_app_js = "\n".join(
        [
            "import { useState } from 'react';",
            "import { FlatList, Pressable, View } from 'react-native';  ①",
            "",
            "export default function App() {",
            "  const [profiles, setProfiles] = useState(initialProfiles);",
            "  const [selectedId, setSelectedId] = useState(initialProfiles[0].id);  ②",
            "",
            "  const selectedProfile =",
            "    profiles.find((profile) => profile.id === selectedId) ?? profiles[0];  ③",
            "",
            "  const renderItem = ({ item }) => (",
            "    <Pressable onPress={() => setSelectedId(item.id)}>  ④",
            "      <View>{/* 목록 카드 */}</View>",
            "    </Pressable>",
            "  );",
            "",
            "  return (",
            "    <>",
            "      <FlatList data={sortedProfiles} renderItem={renderItem} />  ⑤",
            "      {selectedProfile ? (",
            "        <View>{/* 상세 카드 */}</View>  ⑥",
            "      ) : null}",
            "    </>",
            "  );",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "8장. 목록에서 상세 화면으로 이동하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "8.1 왜 상세 화면 이동이 필요한가", 1)
    add_body(document, "목록 화면은 많은 정보를 빠르게 훑어보는 데 적합하지만, 각 항목을 깊게 보여 주기에는 공간이 부족하다. 그래서 실제 앱에서는 목록에서 하나를 선택한 뒤, 그 항목의 자세한 정보를 따로 보여 주는 흐름이 매우 자주 등장한다.")
    add_body(document, "이번 장에서는 복잡한 네비게이션 라이브러리로 넘어가기 전 단계로, 선택 상태를 기반으로 같은 화면 안에 상세 카드를 띄우는 방식을 먼저 익힌다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 선택 상태와 상세 정보 연결")

    add_heading(document, "8.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 지금까지 다룬 목록, 검색, 정렬 구조에 선택 상태 하나만 더하면 목록-상세 연결 흐름을 충분히 연습할 수 있다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "상세 화면 흐름을 처음 배울 때는 실제 페이지 이동부터 붙이기보다, 선택 상태만으로 어떤 데이터가 활성화되었는지 먼저 다루는 편이 구조를 이해하기 훨씬 쉽다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "8.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 08 예제는 목록에서 카드를 누르면, 아래 상세 패널의 이름과 설명이 바뀌는 구조다. 사용자는 지금 어떤 항목을 보고 있는지 시각적으로도 확인할 수 있다.")
    add_command_block(
        document,
        [
            "cd chapters/08_목록에서_상세_화면으로_이동하기/examples/expo-profile-detail",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 다른 카드를 선택해 보자. 선택된 카드 테두리와 하단 상세 카드 내용이 함께 바뀌면 선택 상태 연결이 정상적으로 작동하는 것이다.")

    add_heading(document, "8.4 선택 상태로 상세 정보 바꾸기", 2)
    add_body(document, "이번 예제에서는 목록과 상세가 분리되어 있으면서도 같은 데이터 원본을 공유한다. 사용자가 목록 항목을 누르면 selectedId가 바뀌고, 그 값에 맞는 객체를 찾아 상세 영역에 다시 렌더링한다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "이번 장의 핵심은 Pressable, FlatList, View 조합이다. 목록을 누를 수 있어야 하고, 선택된 정보를 담을 별도 상세 영역도 필요하다."),
            ("②", "selectedId 상태는 현재 사용자가 어떤 항목을 보고 있는지 나타낸다. 상세 화면 흐름의 핵심은 결국 이 활성 선택값을 어떻게 관리하느냐에 달려 있다."),
            ("③", "selectedProfile은 선택된 id에 맞는 실제 데이터 객체다. 상세 패널은 이 객체를 기준으로 텍스트와 상태를 렌더링하게 된다."),
            ("④", "목록 항목을 누를 때 setSelectedId를 호출하면 사용자의 선택이 상태로 기록된다. 이 한 번의 클릭이 상세 패널 내용을 바꾸는 출발점이다."),
            ("⑤", "FlatList는 여전히 목록 UI의 중심이다. 지금까지 배운 검색과 정렬 흐름 위에 선택 기능까지 더해, 목록 화면이 훨씬 실전적으로 확장된다."),
            ("⑥", "상세 카드는 selectedProfile이 있을 때만 보여 준다. 이렇게 하면 선택 상태와 화면 출력이 자연스럽게 연결되고, 화면의 역할도 명확하게 나뉜다."),
        ],
    )
    add_body(document, "이 장에서 중요한 것은 페이지 이동 기법보다, 사용자가 어떤 항목을 선택했고 그 선택이 화면의 어느 부분을 바꾸는지 데이터 흐름으로 이해하는 것이다.")

    document.add_page_break()

    add_heading(document, "8.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 `리액트 쌤` 카드를 선택한 상태다. 선택된 카드는 테두리로 강조되고, 하단 상세 카드에는 현재 프로필의 이름, 역할, 상태 설명이 정리되어 표시된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 8-1. 목록 선택에 따라 상세 정보가 바뀌는 화면", width_cm=7.6)

    add_heading(document, "8.6 정리", 2)
    add_body(document, "이 장에서는 목록에서 항목을 선택하고, 그 결과를 상세 카드에 반영하는 흐름을 완성했다. 독자는 단순 목록 표시를 넘어, 선택과 상세 정보 연결이라는 실제 앱에 가까운 패턴까지 경험할 수 있다.")
    add_body(document, "이제 이 흐름 위에 실제 화면 전환 라이브러리나 탭 구조를 붙이면 더 큰 앱 구조로 확장할 준비가 된 셈이다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 목록의 다른 카드를 눌렀을 때 선택 강조와 하단 상세 카드 내용이 함께 바뀌는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
