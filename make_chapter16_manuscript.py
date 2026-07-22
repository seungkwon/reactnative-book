import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "16_지도_기반_위치_앱_마무리하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "rn-location-hub"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "16_지도_기반_위치_앱_마무리하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "16장 지도 기반 위치 앱 마무리하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "현재 위치와 주소 검색을 하나의 위치 허브 화면으로 통합해 지도 기능 장들을 마무리하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "현재 위치 이동과 주소 검색, 선택 위치 카드, 마커 갱신을 하나의 화면에 묶어 위치 기반 앱의 기본 구조를 완성한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "새 라이브러리를 추가하기보다 14장과 15장의 흐름을 통합하는 데 집중한다. 사용자가 실제 앱처럼 느끼는 화면 구조를 만드는 것이 핵심이다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch16_location_hub.jpg"
    annotated_app_js = "\n".join(
        [
            "const moveMap = (region, title, description) => { ①",
            "  setSelectedPlace({",
            "    title,",
            "    description,",
            "    latitude: region.latitude,",
            "    longitude: region.longitude,",
            "  });",
            "  mapRef.current?.animateToRegion(region, 700);",
            "};",
            "",
            "const moveToCurrentLocation = async () => { ②",
            "  const granted = await requestLocationPermission();",
            "  Geolocation.getCurrentPosition((position) => { ③",
            "    const nextRegion = {",
            "      latitude: position.coords.latitude,",
            "      longitude: position.coords.longitude,",
            "      latitudeDelta: 0.02,",
            "      longitudeDelta: 0.02,",
            "    };",
            "    moveMap(nextRegion, '현재 위치', 'GPS로 가져온 좌표'); ④",
            "  });",
            "};",
            "",
            "const searchAddress = async () => {",
            "  const nextRegion = {",
            "    latitude: result.geometry.location.lat,",
            "    longitude: result.geometry.location.lng,",
            "    latitudeDelta: 0.02,",
            "    longitudeDelta: 0.02,",
            "  };",
            "  moveMap(nextRegion, '검색 결과', result.formatted_address); ⑤",
            "};",
            "",
            "<Marker coordinate={{ latitude: selectedPlace.latitude, longitude: selectedPlace.longitude }} title={selectedPlace.title} description={selectedPlace.description} /> ⑥",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "16장 지도 기반 위치 앱 마무리하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "16.1 위치 앱 흐름을 하나로 묶기", 1)
    add_body(document, "14장에서는 현재 위치를 지도에 표시했고, 15장에서는 주소를 좌표로 바꿔 지도에 반영했다. 이제 남은 과제는 이 두 흐름을 따로 두지 않고, 하나의 사용자 경험으로 묶는 것이다. 실제 앱에서는 사용자가 '내 위치로 가기'와 '주소 검색'을 같은 화면 안에서 오가게 된다.")
    add_body(document, "이번 장은 새 기능을 추가하기보다 기존 기능을 잘 통합하는 데 집중한다. 즉, 현재 위치 버튼과 주소 검색 버튼이 모두 같은 지도 상태를 갱신하고, 선택된 위치 카드가 그 결과를 설명하는 구조로 마무리한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: React Native CLI Android 앱")
    add_bullet(document, "학습 범위: 통합 위치 상태, 현재 위치, 주소 검색, 선택 위치 카드, 지도 마커")

    add_heading(document, "16.2 예제 구성", 2)
    add_body(document, "이번 예제는 react-native-maps와 react-native-geolocation-service를 그대로 사용하지만, 중요한 점은 두 입력 경로를 공통 함수로 묶는 방식이다. 현재 위치로 이동하든, 주소를 검색하든 결국 지도 중심과 선택된 위치 카드, 마커를 함께 갱신해야 하기 때문이다.")
    add_dependency_table(document, package_json)
    add_note_box(
        document,
        "TIP",
        "기능이 늘어날수록 공통으로 바뀌는 상태를 한 함수로 모으는 것이 중요해진다. `moveMap()` 같은 작은 공통 함수를 두면 지도 앱 구조가 훨씬 읽기 쉬워진다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "16.3 현재 위치와 주소 검색을 같은 상태로 연결하기", 1)
    add_body(document, "핵심은 현재 위치와 주소 검색이 서로 다른 출발점이지만 같은 도착점을 가진다는 점이다. 둘 다 결국 어떤 위도, 경도와 설명 문구를 만들고, 그것을 지도 중심과 선택 위치 카드, 마커에 동시에 반영해야 한다.")
    add_command_block(
        document,
        [
            "cd chapters/16_지도_기반_위치_앱_마무리하기/examples/rn-location-hub",
            "npm install",
            "npm run android",
        ],
    )
    add_body(document, "실행 전에 Google Maps API 키와 위치 권한 설정은 앞 장들과 동일하게 필요하다. 이 장은 기능을 통합하는 장이므로 네이티브 준비 조건도 그대로 이어진다.")

    add_heading(document, "16.4 통합 위치 상태 관리", 2)
    add_body(document, "이 장의 중심은 `selectedPlace`와 `moveMap()`이다. 지도 화면에서 바뀌는 정보는 좌표만이 아니라 제목, 설명, 카드 문구, 마커 제목까지 모두 묶여 있으므로, 이를 하나의 상태 덩어리로 관리하면 앱이 훨씬 정돈된다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "`moveMap()`은 이 장의 핵심 공통 함수다. 어떤 입력 경로에서 왔든 지도에 반영할 최종 상태를 한 곳에서 처리해 중복을 줄인다."),
            ("②", "현재 위치 버튼 흐름은 위치 권한과 GPS 좌표 획득으로 시작한다. 하지만 최종 반영은 공통 함수에 맡겨서 구조를 단순하게 유지한다."),
            ("③", "좌표를 실제로 받는 지점이다. 이 콜백 안에서는 새 지역 정보만 만들고, UI 반영 자체는 공통 함수가 담당한다."),
            ("④", "현재 위치 흐름도 결국 제목과 설명이 있는 하나의 위치 객체처럼 다뤄진다. 이 덕분에 카드와 마커가 같은 방식으로 갱신된다."),
            ("⑤", "주소 검색 결과도 현재 위치와 똑같은 `moveMap()` 경로를 타도록 만들면, 기능이 달라도 화면 반응은 일관되게 유지된다."),
            ("⑥", "마커는 `selectedPlace` 상태만 바라본다. 즉, 현재 위치든 검색 결과든 최종 선택 상태만 바꾸면 지도 표시도 자동으로 따라온다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명까지 함께 읽으면, 이번 장은 지도 기능을 더 많이 붙이는 단계가 아니라 이미 만든 기능들을 실제 앱다운 구조로 정돈하는 단계라는 점이 분명해진다.")

    document.add_page_break()

    add_heading(document, "16.5 결과 화면", 1)
    add_body(document, "아래 예시는 주소 입력, 현재 위치 버튼, 검색 버튼, 선택된 위치 카드, 지도 마커가 한 화면에 정리된 통합 위치 앱 모습이다. 이 장까지 마치면 지도 기능 파트의 기본 흐름은 한 사이클을 완성한 셈이다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 16-1. 현재 위치와 주소 검색을 통합한 위치 허브 화면", width_cm=7.6)

    add_heading(document, "16.6 정리", 2)
    add_body(document, "이번 장에서는 현재 위치 이동과 주소 검색, 선택 위치 카드, 마커 갱신을 하나의 위치 허브 화면으로 통합했다. 기능을 추가하는 것보다 상태와 흐름을 정리하는 일이 앱 완성도에 얼마나 중요한지 확인할 수 있다.")
    add_body(document, "이제 지도 파트의 기본기를 마쳤으므로 다음 장에서는 다시 네이티브 SDK 중심 주제인 소셜 로그인과 회원가입으로 넘어가도, 사용자 상태와 외부 플랫폼 응답을 연결하는 감각을 그대로 이어 갈 수 있다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 현재 위치 버튼과 주소 검색 버튼이 모두 같은 카드 상태를 갱신하는지, 지도 중심이 일관되게 이동하는지, 마커 제목과 카드 문구가 함께 바뀌는지 확인해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
