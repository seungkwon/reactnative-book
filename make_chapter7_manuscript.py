import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "07_목록_정렬하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-sort-list"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "07_목록_정렬하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "7장. 목록 정렬하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "정렬 기준에 따라 FlatList 순서를 바꾸기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "정렬 기준 상태를 만들고, 검색 결과를 다시 정렬해 FlatList에 연결하는 방법을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 로컬 배열 정렬에 집중한다. 서버 정렬이나 다중 정렬 옵션의 복잡한 조합은 이후 단계에서 확장할 수 있다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch7_sort_list.jpg"
    annotated_app_js = "\n".join(
        [
            "import { useState } from 'react';",
            "import { FlatList, Pressable } from 'react-native';  ①",
            "",
            "export default function App() {",
            "  const [profiles, setProfiles] = useState(initialProfiles);",
            "  const [query, setQuery] = useState('');",
            "  const [sortKey, setSortKey] = useState('follow');  ②",
            "",
            "  const filteredProfiles = profiles.filter(/* 검색 조건 */);",
            "  const sortedProfiles = [...filteredProfiles].sort((left, right) => {  ③",
            "    if (sortKey === 'name') {",
            "      return left.name.localeCompare(right.name, 'ko');  ④",
            "    }",
            "",
            "    if (left.following === right.following) {",
            "      return left.name.localeCompare(right.name, 'ko');",
            "    }",
            "",
            "    return left.following ? -1 : 1;",
            "  });",
            "",
            "  return (",
            "    <>",
            "      <Pressable onPress={() => setSortKey('follow')} />  ⑤",
            "      <Pressable onPress={() => setSortKey('name')} />",
            "      <FlatList data={sortedProfiles} />  ⑥",
            "    </>",
            "  );",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "7장. 목록 정렬하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "7.1 왜 목록 정렬이 필요한가", 1)
    add_body(document, "검색만으로 원하는 항목을 어느 정도 좁힐 수는 있지만, 결과가 많을 때는 어떤 순서로 보여 줄지도 매우 중요하다. 사용자는 이름순으로 보고 싶을 수도 있고, 중요한 항목을 위에 두고 싶을 수도 있다.")
    add_body(document, "이번 장에서는 chapter 06의 검색 구조를 유지하면서, 정렬 기준에 따라 같은 목록이 서로 다른 순서로 보이도록 확장한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 정렬 기준 상태와 파생 배열")

    add_heading(document, "7.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 외부 정렬 라이브러리 없이 JavaScript의 sort와 상태값 하나만으로도 충분히 직관적인 정렬 기능을 구현할 수 있다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "정렬 기능을 만들 때는 원본 배열을 직접 바꾸기보다, 화면에 보여 줄 복사본을 만들어 정렬하는 편이 흐름을 이해하기 쉽고 부작용도 줄일 수 있다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "7.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 07 예제는 검색 입력을 유지하면서, 팔로잉 우선과 이름순 두 가지 정렬 옵션을 버튼으로 전환할 수 있게 구성했다. 버튼을 바꿀 때마다 같은 데이터라도 목록 순서가 달라진다.")
    add_command_block(
        document,
        [
            "cd chapters/07_목록_정렬하기/examples/expo-sort-list",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 정렬 버튼을 번갈아 눌러 보자. 검색 결과는 유지되면서 카드 순서만 바뀌면 정렬 흐름이 정상적으로 연결된 것이다.")

    add_heading(document, "7.4 정렬 기준으로 순서 바꾸기", 2)
    add_body(document, "이번 예제에서는 먼저 검색으로 결과를 좁힌 뒤, 그 결과를 다시 정렬한다. 즉 검색과 정렬은 경쟁 관계가 아니라, 같은 목록 데이터를 다르게 다루는 두 단계의 처리 흐름으로 이해하는 것이 중요하다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "이번 장의 핵심은 Pressable과 FlatList의 조합이다. 버튼은 정렬 기준을 바꾸고, FlatList는 그 결과 순서를 화면에 반영한다."),
            ("②", "sortKey 상태는 현재 어떤 기준으로 정렬할지 저장한다. 이 값 하나가 바뀌는 것만으로도 같은 목록이 전혀 다른 순서로 보일 수 있다."),
            ("③", "sortedProfiles는 검색 결과를 다시 정렬한 파생 배열이다. 원본 데이터 자체보다 화면에 보여 줄 최종 결과를 계산하는 감각이 중요하다."),
            ("④", "이름순 정렬은 localeCompare를 쓰면 한글 이름도 비교적 자연스럽게 정렬할 수 있다. 단순한 문자열 비교보다 실제 UI에서 더 적절한 선택이다."),
            ("⑤", "정렬 버튼은 상태 변경의 출발점이다. 사용자가 어떤 순서를 원하는지 명시적으로 선택하게 만들면 목록의 제어권을 사용자에게 넘길 수 있다."),
            ("⑥", "FlatList는 더 이상 단순 검색 결과가 아니라 정렬까지 끝난 최종 배열을 받는다. 검색과 정렬이 모두 반영된 화면 결과가 여기서 완성된다."),
        ],
    )
    add_body(document, "정렬 기능의 핵심은 배열을 정렬하는 기술 자체보다, 어떤 시점에 어떤 기준을 적용해 화면용 결과를 계산할지 설계하는 데 있다.")

    document.add_page_break()

    add_heading(document, "7.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 팔로잉 우선 정렬이 선택된 상태다. 이 경우 팔로잉 중인 카드가 먼저 나오고, 같은 그룹 안에서는 이름순으로 정리되어 목록이 보다 읽기 쉬워진다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 7-1. 정렬 기준을 적용한 FlatList 목록 화면", width_cm=7.6)

    add_heading(document, "7.6 정리", 2)
    add_body(document, "이 장에서는 정렬 기준 상태를 두고, 같은 목록 데이터를 서로 다른 순서로 보여 주는 방법을 익혔다. 독자는 검색과 정렬이 결합된 보다 실전적인 목록 UI를 직접 확인할 수 있다.")
    add_body(document, "다음 장에서는 정렬과 검색을 유지한 채, 항목 상세 화면으로 이동하거나 선택 상태를 추가하는 흐름으로 확장하면 인터랙션이 한층 풍부해질 수 있다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 정렬 버튼을 눌렀을 때 같은 검색 결과라도 카드 순서가 달라지는지, 그리고 정렬 기준 전환이 즉시 FlatList에 반영되는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
