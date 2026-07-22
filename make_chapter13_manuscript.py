import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "13_FCM_푸시_메시지_연동하기"
CLIENT_DIR = CHAPTER_DIR / "examples" / "rn-fcm-client"
SERVER_DIR = CHAPTER_DIR / "examples" / "fcm-admin-sender"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "13_FCM_푸시_메시지_연동하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict, title: str) -> None:
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "13장 FCM 푸시 메시지 연동하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "React Native Firebase와 Firebase Admin SDK로 네이티브 푸시 흐름 익히기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "Expo Push가 아닌 FCM 기반의 네이티브 푸시 구조를 이해하고, 앱에서 FCM 토큰을 발급받아 서버에서 테스트 푸시를 보내는 전체 흐름을 정리한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "클라이언트는 React Native Firebase Messaging을 사용하고, 서버는 Firebase Admin SDK를 사용한다. 여기서는 안드로이드 중심의 기본 흐름에 초점을 맞추고, iOS 세부 인증 단계는 소개 수준으로 다룬다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    client_package_json = json.loads(read_text(CLIENT_DIR / "package.json"))
    server_package_json = json.loads(read_text(SERVER_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch13_fcm_push.jpg"
    annotated_client_js = "\n".join(
        [
            "import messaging from '@react-native-firebase/messaging'; ①",
            "",
            "useEffect(() => {",
            "  const unsubscribeForeground = messaging().onMessage(async (remoteMessage) => { ②",
            "    setMessages((current) => [",
            "      { id: `msg-${Date.now()}`, text: remoteMessage.notification?.title ?? '새 메시지' },",
            "      ...current,",
            "    ]);",
            "  });",
            "}, []);",
            "",
            "const setupPush = async () => {",
            "  const authStatus = await messaging().requestPermission(); ③",
            "  await messaging().registerDeviceForRemoteMessages(); ④",
            "  const nextToken = await messaging().getToken(); ⑤",
            "  setToken(nextToken);",
            "};",
            "",
            "messaging().setBackgroundMessageHandler(async (remoteMessage) => { ⑥",
            "  console.log('Background message handled:', remoteMessage);",
            "});",
        ]
    )
    annotated_sender_js = "\n".join(
        [
            "const { initializeApp, applicationDefault } = require('firebase-admin/app');",
            "const { getMessaging } = require('firebase-admin/messaging');",
            "",
            "initializeApp({ credential: applicationDefault() });",
            "",
            "const message = {",
            "  token,",
            "  notification: { title: 'FCM 테스트 메시지', body: 'Firebase Admin SDK에서 보낸 푸시 알림입니다.' },",
            "};",
            "",
            "getMessaging().send(message);",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "13장 FCM 푸시 메시지 연동하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "13.1 왜 Expo Push 대신 FCM을 배워야 할까", 1)
    add_body(document, "푸시 메시지는 앱이 꺼져 있거나 백그라운드에 있어도 사용자를 다시 불러오는 중요한 네이티브 기능이다. Expo Push는 입문용으로 편하지만, 실무에서는 Firebase Cloud Messaging을 직접 다뤄야 하는 경우가 많다. 안드로이드 생태계와의 호환성, 서버 제어 범위, 네이티브 설정 자유도 때문이다.")
    add_body(document, "이번 장에서는 React Native Firebase 공식 Cloud Messaging 문서에서 안내하는 흐름을 기준으로, 앱에서 권한을 요청하고 토큰을 받아 온 뒤, Firebase Admin SDK로 테스트 메시지를 보내는 구조를 정리한다. 여기서 다루는 예제는 2026년 7월 21일 기준 공식 문서의 패키지 이름과 권장 흐름을 따른 것이다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: React Native CLI Android 앱 + Firebase 프로젝트")
    add_bullet(document, "학습 범위: 권한 요청, FCM 토큰 발급, foreground/background 메시지 처리, Admin SDK 발송")

    add_heading(document, "13.2 클라이언트와 서버 패키지 구성", 2)
    add_body(document, "React Native Firebase 공식 문서에 따르면 Messaging 모듈을 쓰기 전에 `@react-native-firebase/app`이 먼저 설정되어 있어야 한다. 메시지 수신은 `@react-native-firebase/messaging`에서 담당한다. 서버 쪽은 Firebase 공식 서버 환경 문서가 권장하는 Firebase Admin SDK를 사용한다.")
    add_dependency_table(document, client_package_json, "클라이언트 패키지")
    add_body(document, "")
    add_dependency_table(document, server_package_json, "서버 패키지")

    add_note_box(
        document,
        "TIP",
        "FCM은 단순히 라이브러리 설치만으로 끝나지 않는다. 실제 프로젝트에서는 Firebase Console 설정, `google-services.json`, Android Manifest, 알림 채널 같은 네이티브 설정도 함께 맞춰야 한다. 이번 장 예제는 그중 자바스크립트 흐름을 이해하는 데 집중한다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "13.3 권한 요청과 토큰 발급", 1)
    add_body(document, "공식 Cloud Messaging 문서에 따르면 iOS에서는 권한 요청이 필수이고, React Native 앱에서는 `requestPermission()`으로 권한 상태를 확인할 수 있다. 또한 `onMessage()`는 포그라운드 상태 메시지를 처리할 때 사용하고, `setBackgroundMessageHandler()`는 백그라운드나 종료 상태 메시지 처리를 위해 앱 시작 지점에서 가능한 빨리 등록해야 한다.")
    add_command_block(
        document,
        [
            "cd chapters/13_FCM_푸시_메시지_연동하기/examples/rn-fcm-client",
            "npm install",
            "npm run android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 권한 요청 버튼을 누르면 권한 상태가 갱신되고, 토큰 문자열이 화면에 표시된다. 이 토큰을 복사해 서버 예제에 넘기면 테스트 푸시를 바로 보낼 수 있다.")

    add_heading(document, "13.4 메시지 수신 흐름을 코드로 연결하기", 2)
    add_body(document, "푸시 기능은 크게 세 부분으로 나눠 생각하면 쉽다. 첫째는 권한과 토큰 발급, 둘째는 포그라운드 수신 처리, 셋째는 백그라운드 수신 처리다. 공식 문서를 기준으로 이 세 흐름을 화면 상태와 연결하면 푸시 기능의 뼈대가 완성된다.")
    add_code_block(document, "App.js / index.js", annotated_client_js)
    add_code_notes(
        document,
        [
            ("①", "공식 문서 기준 메시지 기능은 `@react-native-firebase/messaging` 모듈에서 가져온다. 이 모듈은 `@react-native-firebase/app`이 먼저 준비되어 있어야 정상 동작한다."),
            ("②", "포그라운드 상태에서 도착한 메시지는 `onMessage()`로 받는다. 이 핸들러에서는 Alert를 띄우거나 화면 상태를 바꾸는 등 즉시 UI 반응을 줄 수 있다."),
            ("③", "권한 요청은 푸시 기능의 시작점이다. 특히 iOS에서는 명시적 허용이 필요하고, 안드로이드에서도 사용자 경험상 권한 상태를 화면에서 확인해 주는 편이 좋다."),
            ("④", "공식 문서에는 자동 등록을 끄는 경우 `registerDeviceForRemoteMessages()`를 앱 시작 초기에 호출하라고 안내한다. 예제에서는 흐름을 이해하기 쉽게 권한 요청 직후에 함께 배치했다."),
            ("⑤", "`getToken()`으로 받은 FCM 토큰은 서버가 특정 디바이스를 대상으로 푸시를 보낼 때 사용한다. 실무에서는 이 값을 로그인 사용자와 함께 서버 DB에 저장하는 경우가 많다."),
            ("⑥", "백그라운드 핸들러는 앱 컴포넌트 바깥, 가능한 이른 시점에 등록해야 한다. 공식 문서도 `index.js` 같은 진입점에서 `setBackgroundMessageHandler()`를 설정하는 방식을 권장한다."),
        ],
    )
    add_body(document, "서버 발송 코드는 매우 짧지만 의미는 크다. Firebase 공식 서버 문서는 서버 환경에서 Firebase Admin SDK 사용을 권장하며, `getMessaging().send(message)`로 특정 토큰에 직접 푸시를 전송할 수 있다고 설명한다.")
    add_code_block(document, "send-message.js", annotated_sender_js)

    document.add_page_break()

    add_heading(document, "13.5 결과 화면", 1)
    add_body(document, "아래 예시는 왼쪽에서 앱이 토큰을 발급받고 수신 로그를 표시하는 모습, 오른쪽에서 Admin SDK 스크립트가 메시지를 성공적으로 보낸 콘솔 결과를 함께 정리한 것이다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 13-1. FCM 토큰 발급과 Admin SDK 발송 예시", width_cm=12.2)

    add_heading(document, "13.6 정리", 2)
    add_body(document, "이번 장에서는 Expo Push 대신 FCM 기반의 네이티브 푸시 흐름을 React Native와 Firebase 공식 문서 기준으로 정리했다. 앱에서는 권한 요청과 토큰 발급, 포그라운드/백그라운드 수신 처리 구조를 만들었고, 서버에서는 Firebase Admin SDK로 테스트 푸시를 보내는 예제를 준비했다.")
    add_body(document, "다음 장에서 지도와 위치 기능으로 넘어가더라도, 이번 장에서 익힌 핵심은 그대로 남는다. 네이티브 기능은 결국 권한 요청, 식별자 또는 상태 획득, 그리고 서버나 플랫폼 이벤트와의 연결이라는 같은 패턴으로 이해할 수 있다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 권한 요청 후 상태가 바뀌는지, 토큰이 화면에 표시되는지, 포그라운드 메시지가 로그 카드에 추가되는지, Admin SDK가 메시지 ID를 반환하는지 차례대로 점검해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
