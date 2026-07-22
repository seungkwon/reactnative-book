import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "15_주소를_좌표로_변환해_지도에_표시하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "rn-address-geocode-map"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "15_주소를_좌표로_변환해_지도에_표시하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "15장 주소를 좌표로 변환해 지도에 표시하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "주소 검색 문자열을 Geocoding API로 좌표화하고 지도 중심과 마커를 함께 갱신하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "사용자가 입력한 주소를 Geocoding API로 좌표로 변환하고, 그 결과를 지도 중심과 마커, 상태 카드에 반영한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "이번 장은 Google Geocoding API를 기준으로 주소를 좌표로 바꾸는 흐름에 집중한다. 역지오코딩과 경로 탐색, 장소 자동완성은 다음 확장 단계로 남겨 둔다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch15_address_geocode.jpg"
    annotated_app_js = "\n".join(
        [
            "const GOOGLE_MAPS_API_KEY = 'YOUR_GOOGLE_MAPS_API_KEY'; ①",
            "",
            "const searchAddress = async () => {",
            "  const encodedAddress = encodeURIComponent(address.trim()); ②",
            "  const response = await fetch(",
            "    `https://maps.googleapis.com/maps/api/geocode/json?address=${encodedAddress}&key=${GOOGLE_MAPS_API_KEY}` ③",
            "  );",
            "  const data = await response.json();",
            "",
            "  const result = data.results[0]; ④",
            "  const location = result.geometry.location;",
            "  const nextRegion = {",
            "    latitude: location.lat,",
            "    longitude: location.lng,",
            "    latitudeDelta: 0.02,",
            "    longitudeDelta: 0.02,",
            "  };",
            "",
            "  setSearchResult({ latitude: location.lat, longitude: location.lng, title: result.formatted_address }); ⑤",
            "  mapRef.current?.animateToRegion(nextRegion, 700); ⑥",
            "};",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "15장 주소를 좌표로 변환해 지도에 표시하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "15.1 주소 검색과 지도의 연결", 1)
    add_body(document, "현재 위치 지도까지 만들었다면 다음으로 자연스럽게 이어지는 기능은 주소 검색이다. 사용자가 '서울시청'이나 '세종대로 110' 같은 문자열을 입력하면, 앱은 그것을 실제 좌표로 바꿔 지도 위 위치로 보여줘야 한다. 이때 필요한 것이 Geocoding API다.")
    add_body(document, "Google Geocoding API 공식 문서는 주소를 좌표로 바꾸는 과정을 geocoding 또는 forward geocoding이라고 설명한다. 공식 예시에는 주소를 포함한 HTTP 요청을 보내고, 응답의 latitude와 longitude를 읽어 지도 중심이나 마커 위치에 사용할 수 있다고 안내한다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: React Native CLI Android 앱")
    add_bullet(document, "학습 범위: 주소 입력, Geocoding API 요청, 지도 중심 이동, 검색 결과 마커")

    add_heading(document, "15.2 예제 구성과 API 키 주의사항", 2)
    add_body(document, "이번 예제는 지도 렌더링에는 react-native-maps를 사용하고, 주소 검색은 fetch로 Google Geocoding API를 직접 호출한다. 구조는 단순하지만 실무에서는 API 키 노출을 막기 위한 제한 정책이나 서버 프록시를 함께 고려해야 한다.")
    add_dependency_table(document, package_json)
    add_note_box(
        document,
        "TIP",
        "Google 공식 문서 기준으로 주소 검색은 `https://maps.googleapis.com/maps/api/geocode/json?address=...&key=...` 형태의 요청으로도 이해할 수 있다. 책 예제에서는 흐름을 쉽게 설명하기 위해 이 구조를 사용한다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "15.3 주소를 좌표로 바꿔 지도에 반영하기", 1)
    add_body(document, "검색 버튼을 누르면 입력한 주소를 인코딩하고 Geocoding API를 호출한다. 응답에서 첫 번째 결과를 꺼낸 뒤 위도와 경도를 읽어 지도 중심과 마커 상태를 함께 갱신하면, 사용자는 입력한 주소가 실제 지도 위치로 어떻게 바뀌는지 한눈에 볼 수 있다.")
    add_command_block(
        document,
        [
            "cd chapters/15_주소를_좌표로_변환해_지도에_표시하기/examples/rn-address-geocode-map",
            "npm install",
            "npm run android",
        ],
    )
    add_body(document, "실행 전에 Geocoding API가 활성화된 Google Maps API 키를 준비해야 한다. 또한 지도 렌더링을 위해 react-native-maps에서 사용하는 지도 API 키 설정도 이미 되어 있어야 한다.")

    add_heading(document, "15.4 Geocoding 응답을 지도 상태로 바꾸기", 2)
    add_body(document, "이 장의 핵심은 텍스트 주소를 API 결과로 바꾸고, 다시 그 결과를 지도 상태로 변환하는 과정이다. 입력 문자열, 네트워크 응답, 지도 UI가 하나의 흐름으로 연결되기 때문에 주소 검색 기능이 완성된다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "API 키는 실제 서비스에서는 환경 변수나 서버 프록시 뒤에 감추는 편이 안전하다. 예제에서는 구조를 설명하기 위해 상수 자리에 두었다."),
            ("②", "사용자가 입력한 주소에는 공백이나 특수문자가 포함될 수 있으므로 요청 전에 `encodeURIComponent()`로 안전하게 변환해야 한다."),
            ("③", "공식 Geocoding 문서의 핵심은 주소를 포함한 HTTP 요청을 보내고 좌표 응답을 받는 구조다. 이 URL이 바로 forward geocoding의 기본 형태다."),
            ("④", "응답 결과가 여러 개일 수 있지만 입문 예제에서는 첫 번째 결과를 사용하면 흐름을 이해하기 쉽다. 결과가 없을 때는 상태 문구로 실패 이유를 보여 주는 편이 좋다."),
            ("⑤", "응답에서 꺼낸 위도, 경도, 주소명을 상태에 저장하면 마커 제목과 좌표 카드가 함께 갱신된다. 하나의 API 결과를 여러 UI 요소에 재사용하는 전형적인 패턴이다."),
            ("⑥", "지도 중심도 같은 좌표로 이동시켜야 검색 결과를 시각적으로 바로 확인할 수 있다. animateToRegion은 이 연결을 자연스럽게 만들어 준다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명까지 함께 읽으면, 주소 검색은 단순한 입력창 기능이 아니라 외부 지도 플랫폼의 응답을 앱 상태와 화면 경험으로 바꾸는 전형적인 네이티브 연동 예제라는 점이 분명해진다.")

    document.add_page_break()

    add_heading(document, "15.5 결과 화면", 1)
    add_body(document, "아래 예시는 주소 입력창, 검색 버튼, 상태 문구, 검색 결과 카드, 지도 마커가 한 화면에 정리된 모습이다. 다음 장에서는 이런 지도 기능을 조금 더 통합된 위치 앱 구조로 마무리한다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 15-1. 주소 검색 결과를 지도에 표시한 화면", width_cm=7.6)

    add_heading(document, "15.6 정리", 2)
    add_body(document, "이번 장에서는 주소 문자열을 Geocoding API로 좌표로 변환하고, 그 결과를 지도 중심과 마커, 상태 카드에 반영하는 흐름을 만들었다. 지도 기능이 단순 렌더링을 넘어 검색과 외부 API 연동으로 확장되는 첫 단계다.")
    add_body(document, "다음 장에서는 현재 위치, 주소 검색, 마커 갱신 같은 조각 기능을 하나의 위치 기반 앱 흐름으로 묶어 마무리한다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 주소 입력 후 상태 문구가 바뀌는지, 좌표 카드가 갱신되는지, 지도 중심이 이동하는지, 마커 제목이 검색 주소로 바뀌는지 확인해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
