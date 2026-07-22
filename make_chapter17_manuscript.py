import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "17_Native_SDK로_소셜_로그인과_회원가입_구현하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "rn-social-auth"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "17_Native_SDK로_소셜_로그인과_회원가입_구현하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 설명", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "17장 Native SDK로 소셜 로그인과 회원가입 구현하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "카카오 Native SDK와 Apple 로그인으로 회원가입 시작 화면을 완성하기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "카카오 로그인과 Apple 로그인을 React Native 회원가입 화면에 연결하고, 로그인 결과를 가입 완료 상태 카드로 이어 붙인다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "학습 범위",
        "이번 장은 SDK 연동 흐름과 화면 구조에 집중한다. 실제 서비스에서는 로그인 성공 후 받은 토큰을 서버에서 검증해 자체 회원 체계와 연결해야 한다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch17_social_auth.jpg"
    annotated_app_js = "\n".join(
        [
            "import appleAuth, { AppleButton } from '@invertase/react-native-apple-authentication'; ①",
            "import { login, me } from '@react-native-kakao/user'; ②",
            "",
            "const signInWithKakao = async () => {",
            "  await login(); ③",
            "  const user = await me();",
            "  setMemberProfile({ provider: 'kakao', name: user.nickname, email: user.email }); ④",
            "};",
            "",
            "const signInWithApple = async () => {",
            "  const response = await appleAuth.performRequest({ ⑤",
            "    requestedOperation: appleAuth.Operation.LOGIN,",
            "    requestedScopes: [appleAuth.Scope.FULL_NAME, appleAuth.Scope.EMAIL],",
            "  });",
            "  setMemberProfile({ provider: 'apple', name: response.fullName?.givenName, email: response.email });",
            "};",
            "",
            "<Text style={styles.profileValue}>이메일: {memberProfile?.email ?? '로그인 후 표시'}</Text> ⑥",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "17장 Native SDK로 소셜 로그인과 회원가입 구현하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "17.1 왜 Native SDK 소셜 로그인을 따로 배워야 할까", 1)
    add_body(document, "소셜 로그인은 단순히 버튼 하나를 붙이는 문제가 아니다. 각 플랫폼이 제공하는 인증 절차, 네이티브 앱 전환, 동의 화면, 토큰 반환 방식이 얽혀 있기 때문에 웹 OAuth만 알던 감각으로 접근하면 흐름이 쉽게 끊긴다.")
    add_body(document, "2026년 7월 22일 기준 카카오 공식 문서는 Android SDK에서 `loginWithKakaoTalk()`과 `loginWithKakaoAccount()`를 모두 지원한다고 안내하고 있으며, Apple 공식 문서는 iOS 앱에서 Sign in with Apple capability와 Authentication Services 기반 흐름을 요구한다. 따라서 이번 장은 각 플랫폼 고유 인증 흐름을 React Native 화면에 연결하는 감각을 익히는 데 초점을 둔다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: React Native CLI 앱, Android 카카오 로그인, iOS Apple 로그인")
    add_bullet(document, "학습 범위: 카카오 로그인, Apple 로그인, 회원가입 완료 카드, 서버 검증 전 단계")

    add_heading(document, "17.2 예제 패키지와 사전 설정", 2)
    add_body(document, "카카오 로그인은 공식 Kakao SDK 흐름을 감싸는 React Native 래퍼를 사용해 예제를 구성하고, Apple 로그인은 React Native용 Apple Authentication 라이브러리를 이용해 iOS Authentication Services 흐름을 호출한다. Apple 공식 문서에 따르면 App ID에 Sign in with Apple capability를 켜야 하며, 카카오 공식 문서에 따르면 카카오 로그인을 사용하려면 네이티브 앱 키와 플랫폼 등록, 리다이렉트 설정이 필요하다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "카카오와 Apple 모두 로그인 성공만으로 가입이 끝나는 것은 아니다. 실제 서비스에서는 SDK가 돌려준 access token이나 identity token을 서버에 보내 검증한 뒤, 자체 회원 테이블에 연결해야 안전하다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "17.3 소셜 로그인 결과를 회원가입 화면으로 연결하기", 1)
    add_body(document, "이번 예제는 로그인 성공 직후 프로필 카드를 갱신해 회원가입 완료 흐름을 시각적으로 보여 준다. 즉, SDK 호출과 토큰 반환 자체보다 그 결과를 앱 가입 상태와 어떻게 연결하는지가 중심이다.")
    add_command_block(
        document,
        [
            "cd chapters/17_Native_SDK로_소셜_로그인과_회원가입_구현하기/examples/rn-social-auth",
            "npm install",
            "npm run android",
            "npm run ios",
        ],
    )
    add_body(document, "Android에서는 카카오 로그인 흐름을, iOS에서는 Apple 로그인 흐름을 각각 테스트하면 된다. Apple 공식 문서상 Sign in with Apple은 iOS 네이티브 capability와 팀 설정이 먼저 준비되어 있어야 정상 동작한다.")

    add_heading(document, "17.4 카카오와 Apple 로그인 상태를 한 화면에 묶기", 2)
    add_body(document, "이 장의 핵심은 서로 다른 로그인 SDK 결과를 하나의 `memberProfile` 상태로 모으는 것이다. 이렇게 해야 어떤 소셜 계정으로 들어왔든 가입 완료 카드와 후속 회원가입 흐름이 같은 구조로 이어질 수 있다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "Apple 로그인은 iOS 네이티브의 Authentication Services 흐름을 호출하는 라이브러리를 통해 사용한다. Apple 공식 문서도 이 프레임워크 기반 구현을 안내한다."),
            ("②", "카카오 로그인은 공식 Kakao SDK 흐름을 React Native에서 다루기 쉽게 연결해 주는 모듈을 사용한다. 핵심은 카카오톡 로그인 또는 카카오계정 로그인의 결과를 JS 상태로 끌어오는 것이다."),
            ("③", "카카오 공식 Android 문서 기준 권장 흐름은 카카오톡이 가능하면 카카오톡으로 로그인하고, 그렇지 않으면 계정 로그인으로 넘어가는 방식이다. 예제에서는 그 시작점을 단일 함수로 다뤘다."),
            ("④", "로그인 성공 후 `me()` 같은 사용자 정보 조회 결과를 가입 상태 카드에 바로 연결하면, 사용자는 단순 로그인 성공을 넘어 회원가입 완료로 인식하게 된다."),
            ("⑤", "Apple 공식 문서에 따르면 로그인 요청 시 이름과 이메일 같은 스코프를 요청할 수 있다. 다만 Apple은 최초 동의 시점에만 일부 정보를 제공할 수 있으므로 서버 저장 전략이 중요하다."),
            ("⑥", "결국 중요한 것은 어떤 공급자(provider)로 들어왔든 결과를 같은 UI 상태로 정리하는 것이다. 이 패턴이 있어야 이후 약관 동의나 추가 정보 입력 단계도 같은 방식으로 확장할 수 있다."),
        ],
    )
    add_body(document, "코드 스니펫 아래 설명까지 함께 읽으면, 이번 장은 소셜 로그인 SDK 자체를 외우는 장이 아니라 외부 인증 결과를 앱 회원 상태로 연결하는 구조를 배우는 장이라는 점이 더 분명해진다.")

    document.add_page_break()

    add_heading(document, "17.5 결과 화면", 1)
    add_body(document, "아래 예시는 카카오 버튼, Apple 버튼, 상태 문구, 가입 완료 카드가 한 화면에 정리된 모습이다. 다음 단계에서는 이 가입 완료 상태를 실제 서버 회원 정보와 연결하면 완전한 로그인/회원가입 기능으로 확장된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 17-1. 카카오와 Apple 로그인을 연결한 회원가입 화면", width_cm=7.6)

    add_heading(document, "17.6 정리", 2)
    add_body(document, "이번 장에서는 카카오 Native SDK 기반 로그인 흐름과 Apple 로그인 흐름을 React Native 화면에 연결하고, 그 결과를 하나의 회원가입 완료 카드로 정리했다. 공급자는 달라도 로그인 결과를 같은 회원 상태로 연결하는 패턴이 핵심이다.")
    add_body(document, "다음 장에서는 네이티브 앱과 웹 화면을 동시에 다루는 WebView 하이브리드 구조를 통해, 또 다른 형태의 플랫폼 연동 방식을 정리한다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 부분은 네 가지다. 카카오 로그인 성공 후 카드가 갱신되는지, Apple 로그인 버튼이 iOS에서만 동작하는지, 상태 문구가 공급자에 맞게 바뀌는지, 서버 검증 필요성을 코드와 원고에서 함께 이해했는지 점검해 보자.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
