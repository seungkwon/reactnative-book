import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from make_it_book_template import (
    add_header_footer,
    add_note_box,
    add_run,
    configure_page,
    make_styles,
    set_cell_shading,
    set_paragraph_border,
    set_paragraph_shading,
)


ROOT = Path(__file__).resolve().parent
CHAPTER_DIR = ROOT / "chapters" / "05_입력으로_목록_항목_추가하기"
EXAMPLE_DIR = CHAPTER_DIR / "examples" / "expo-add-list-item"
ARTIFACTS_DIR = CHAPTER_DIR / "artifacts"
MANUSCRIPT_PATH = CHAPTER_DIR / "manuscript" / "05_입력으로_목록_항목_추가하기.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_run(p, text, size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    add_run(
        p,
        text,
        bold=True,
        color="1A365D" if level == 1 else "B85C38",
        size=18 if level == 1 else 14,
    )


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    add_run(p, "- ", bold=True, color="1A365D", size=11)
    add_run(p, text, size=11)


def add_code_block(document: Document, title: str, code: str) -> None:
    markers = {"①", "②", "③", "④", "⑤", "⑥"}
    p = document.add_paragraph()
    add_run(p, title, bold=True, color="1A365D", size=11)
    for line in code.splitlines():
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F8FAFC")
        set_paragraph_border(cp, color="CBD5E1", size=8, space=2)
        if not line:
            add_run(cp, " ", font="Consolas", size=9.5)
            continue
        chunks = line.split()
        if chunks and chunks[-1] in markers:
            marker = chunks[-1]
            code_text = line[: line.rfind(marker)].rstrip()
            add_run(cp, code_text + "  ", font="Consolas", size=9.5)
            add_run(cp, marker, font="맑은 고딕", size=11.5, bold=True, color="2563EB")
        else:
            add_run(cp, line, font="Consolas", size=9.5)


def add_code_notes(document: Document, notes: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    add_run(p, "코드 스니펫 해설", bold=True, color="1A365D", size=11)
    for marker, text in notes:
        note = document.add_paragraph()
        note.paragraph_format.left_indent = Cm(0.4)
        note.paragraph_format.space_after = Pt(4)
        add_run(note, f"{marker} ", bold=True, color="1D4ED8", size=11)
        add_run(note, text, size=11)


def add_command_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    add_run(p, "실행 명령", bold=True, color="1A365D", size=11)
    for line in lines:
        cp = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(cp, "F9FAFB")
        set_paragraph_border(cp, color="D1D5DB", size=8, space=2)
        add_run(cp, line, font="Consolas", size=9.5)


def add_dependency_table(document: Document, package_json: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, text in enumerate(["패키지", "버전"]):
        set_cell_shading(header[idx], "DCE6F1")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hp = header[idx].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, text, bold=True, color="1A365D", size=10.5)

    for name, version in package_json["dependencies"].items():
        row = table.add_row().cells
        add_run(row[0].paragraphs[0], name, size=10.5)
        add_run(row[1].paragraphs[0], version, size=10.5)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 8.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color="555555", size=10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    add_run(p, "5장. 입력으로 목록 항목 추가하기", bold=True, color="1A365D", size=24)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "TextInput과 버튼으로 FlatList 데이터를 늘리기", color="555555", size=13)

    add_note_box(
        document,
        "학습 목표",
        "TextInput 입력값을 이용해 새 데이터를 만들고, 버튼 클릭으로 FlatList 목록에 새 항목을 추가하는 방법을 이해한다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "실습 범위",
        "이 장은 목록 추가 기능에 집중한다. 수정, 삭제, 서버 저장 같은 기능은 뒤 장에서 다룰 수 있도록 남겨 둔다.",
        "FFF7ED",
        "C2410C",
    )


def main() -> None:
    package_json = json.loads(read_text(EXAMPLE_DIR / "package.json"))
    ui_screen = ARTIFACTS_DIR / "ch5_add_list_item.jpg"
    annotated_app_js = "\n".join(
        [
            "import { StatusBar } from 'expo-status-bar';",
            "import { useState } from 'react';",
            "import { FlatList, Image, Pressable, SafeAreaView, StyleSheet, Text, TextInput, View } from 'react-native';  ①",
            "",
            "export default function App() {",
            "  const [profiles, setProfiles] = useState(initialProfiles);",
            "  const [nickname, setNickname] = useState('');  ②",
            "",
            "  const addProfile = () => {",
            "    const trimmedName = nickname.trim();  ③",
            "    if (!trimmedName) {",
            "      return;",
            "    }",
            "",
            "    const newProfile = {",
            "      id: String(Date.now()),",
            "      name: trimmedName,  ④",
            "      role: 'New FlatList Member',",
            "      following: false,",
            "    };",
            "",
            "    setProfiles((currentProfiles) => [newProfile, ...currentProfiles]);  ⑤",
            "    setNickname('');",
            "  };",
            "",
            "  return (",
            "    <SafeAreaView style={styles.screen}>",
            "      <TextInput value={nickname} onChangeText={setNickname} />",
            "      <Pressable onPress={addProfile}>  ⑥",
            "        <Text>추가</Text>",
            "      </Pressable>",
            "      <FlatList data={profiles} />",
            "    </SafeAreaView>",
            "  );",
            "}",
        ]
    )

    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "5장. 입력으로 목록 항목 추가하기")

    for section in document.sections:
        configure_page(section)

    add_cover(document)

    add_heading(document, "5.1 왜 목록 추가 기능이 필요한가", 1)
    add_body(document, "chapter 04에서는 준비된 배열 데이터를 FlatList로 화면에 보여 주는 데 집중했다. 하지만 실제 앱에서는 사용자가 직접 새로운 데이터를 만들어 목록에 넣는 경우가 많다.")
    add_body(document, "메모 앱의 새 할 일, 연락처 앱의 새 사람, 커뮤니티 앱의 새 항목처럼 입력과 목록 추가는 매우 자주 등장하는 조합이다. 이번 장에서는 그 가장 기본적인 패턴을 작은 예제로 익힌다.")
    add_bullet(document, "운영체제: Windows")
    add_bullet(document, "실행 대상: Android 에뮬레이터")
    add_bullet(document, "학습 범위: 입력 상태와 목록 추가")

    add_heading(document, "5.2 준비 환경과 패키지", 2)
    add_body(document, "이번 예제도 Expo 기본 프로젝트 위에서 진행한다. 입력값을 저장하는 상태와 목록 배열 상태를 함께 관리하는 구조만으로도 충분히 실용적인 예제를 만들 수 있다.")
    add_dependency_table(document, package_json)

    add_note_box(
        document,
        "TIP",
        "목록 추가 기능을 만들 때는 먼저 입력값을 비우는 타이밍과, 새 항목을 배열 앞이나 뒤 중 어디에 넣을지를 분명히 정해 두면 흐름이 훨씬 명확해진다.",
        "F0FDF4",
        "15803D",
    )

    document.add_page_break()

    add_heading(document, "5.3 예제 프로젝트 실행", 1)
    add_body(document, "chapter 05 예제는 입력창에 이름을 적고 추가 버튼을 누르면, 그 이름이 새 프로필 카드가 되어 목록의 맨 앞에 나타난다. 사용자는 자신의 입력이 목록 구조를 직접 바꾸는 장면을 바로 확인할 수 있다.")
    add_command_block(
        document,
        [
            "cd chapters/05_입력으로_목록_항목_추가하기/examples/expo-add-list-item",
            "npm install",
            "npx expo start --android",
        ],
    )
    add_body(document, "앱을 실행한 뒤 입력창에 임의의 이름을 입력하고 추가 버튼을 눌러 보자. 새 카드가 곧바로 맨 위에 들어오면 상태 갱신과 목록 재렌더링이 정상적으로 이어진 것이다.")

    add_heading(document, "5.4 입력값으로 새 항목 만들기", 2)
    add_body(document, "이번 예제의 핵심은 입력 상태와 목록 상태를 함께 다루는 것이다. 사용자가 입력창에 적은 문자열을 바탕으로 새 객체를 만들고, 이를 기존 배열 앞쪽에 붙인 뒤, FlatList가 다시 그려지도록 만든다.")
    add_code_block(document, "App.js", annotated_app_js)
    add_code_notes(
        document,
        [
            ("①", "이번 장에서는 기존 FlatList 구조에 TextInput이 추가된다. 이제 화면은 단순히 데이터를 보여 주는 수준을 넘어, 사용자의 입력을 받아 새로운 데이터를 만들어 낼 수 있게 된다."),
            ("②", "nickname 상태는 입력창의 현재 값을 저장한다. 목록 상태와 별도로 입력 상태를 분리해 두면, 사용자가 무엇을 입력 중인지와 이미 저장된 목록 데이터를 명확히 나눠 관리할 수 있다."),
            ("③", "trim을 거치면 공백만 입력한 경우를 걸러낼 수 있다. 작은 검증이지만, 실제 목록 추가 기능에서는 이런 기본 방어가 꽤 중요하다."),
            ("④", "새 항목 객체를 만들 때는 화면에 필요한 속성을 한 번에 정의해 두는 것이 좋다. name, role, following처럼 카드 렌더링에 필요한 구조를 맞춰 두면 기존 목록과 같은 컴포넌트로 바로 출력할 수 있다."),
            ("⑤", "새 항목을 배열 앞에 붙이면 방금 추가한 카드가 즉시 눈에 띄게 된다. 학습용 예제에서는 이 방식이 결과를 가장 분명하게 보여 준다."),
            ("⑥", "추가 버튼은 입력 상태와 목록 상태를 연결하는 출발점이다. 사용자의 클릭 한 번이 새 객체 생성, 배열 갱신, FlatList 재렌더링으로 이어지는 흐름을 꼭 함께 이해해야 한다."),
        ],
    )
    add_body(document, "이 장을 통해 독자는 React Native 화면이 단순히 준비된 데이터를 보여 주는 공간이 아니라, 사용자의 입력을 받아 데이터를 늘리고 다시 렌더링하는 인터랙티브한 구조라는 점을 더 분명하게 체감할 수 있다.")

    document.add_page_break()

    add_heading(document, "5.5 실행 결과 확인", 1)
    add_body(document, "아래 화면은 입력창에 `새 프로필`을 적고 추가 버튼을 눌렀을 때의 결과다. 새 카드가 목록 최상단에 들어오고, 기존 카드들은 아래로 밀리면서 그대로 유지된다.")
    if ui_screen.exists():
        add_image(document, ui_screen, "그림 5-1. 입력값으로 새 항목을 추가한 FlatList 화면", width_cm=7.6)

    add_heading(document, "5.6 정리", 2)
    add_body(document, "이 장에서는 TextInput과 버튼을 이용해 FlatList 목록에 새 항목을 추가했다. 독자는 입력 상태, 배열 상태, 재렌더링이 어떻게 연결되는지 실제 동작으로 확인할 수 있다.")
    add_body(document, "다음 장에서는 추가된 항목을 검색하거나 조건에 따라 걸러내는 방식으로 확장하면, 목록 데이터 다루기가 한층 더 실전적인 흐름으로 이어진다.")
    add_note_box(
        document,
        "체크포인트",
        "독자가 직접 확인해야 할 핵심은 이름을 입력한 뒤 추가 버튼을 눌렀을 때 새 카드가 목록 맨 앞에 생성되는지, 그리고 입력창이 다시 비워지는지 여부다.",
        "FEF2F2",
        "B91C1C",
    )

    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(MANUSCRIPT_PATH))
    print(str(MANUSCRIPT_PATH))


if __name__ == "__main__":
    main()
