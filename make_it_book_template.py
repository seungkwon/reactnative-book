from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_NAME = "IT_책_집필_워드_템플릿_ko.docx"


def set_east_asia_font(target, font_name):
    rpr = target._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_paragraph_border(paragraph, color="D9D9D9", size=8, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        el = p_bdr.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            p_bdr.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def add_run(paragraph, text, *, bold=False, italic=False, color=None, font="맑은 고딕", size=11):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    set_east_asia_font(run, font)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def configure_page(section):
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def make_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(11)
    set_east_asia_font(normal, "맑은 고딕")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    title = document.styles["Title"]
    title.font.name = "맑은 고딕"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    set_east_asia_font(title, "맑은 고딕")

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "맑은 고딕"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_east_asia_font(subtitle, "맑은 고딕")

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "맑은 고딕"
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    set_east_asia_font(heading1, "맑은 고딕")

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "맑은 고딕"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0xB8, 0x5C, 0x38)
    set_east_asia_font(heading2, "맑은 고딕")

    styles = document.styles

    if "본문 강조" not in styles:
        emph = styles.add_style("본문 강조", WD_STYLE_TYPE.PARAGRAPH)
        emph.base_style = styles["Normal"]
        emph.font.name = "맑은 고딕"
        emph.font.size = Pt(11)
        emph.font.bold = True
        emph.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        set_east_asia_font(emph, "맑은 고딕")

    if "코드 블록" not in styles:
        code = styles.add_style("코드 블록", WD_STYLE_TYPE.PARAGRAPH)
        code.base_style = styles["Normal"]
        code.font.name = "Consolas"
        code.font.size = Pt(9.5)
        set_east_asia_font(code, "맑은 고딕")
        code.paragraph_format.left_indent = Cm(0.4)
        code.paragraph_format.right_indent = Cm(0.2)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.line_spacing = 1.1


def add_header_footer(section, header_text):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, header_text, color="6B7280", size=9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Python으로 쓰는 IT 책 템플릿  |  ", color="6B7280", size=9)
    add_page_number(footer)


def add_note_box(document, title, body, fill, border_color):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=border_color, size=10, space=3)
    add_run(p, f"{title}  ", bold=True, color=border_color, size=10.5)
    add_run(p, body, size=10.5)


def add_code_line(paragraph, indent, parts):
    if indent:
        add_run(paragraph, " " * indent, font="Consolas", size=9.5)
    for text, color, bold in parts:
        add_run(paragraph, text, font="Consolas", size=9.5, color=color, bold=bold)


def add_code_block(document):
    lines = [
        (0, [("import", "8B5CF6", True), (" ", None, False), ("time", "111827", False)]),
        (0, [("from", "8B5CF6", True), (" ", None, False), ("pathlib", "111827", False), (" ", None, False), ("import", "8B5CF6", True), (" ", None, False), ("Path", "0F766E", False)]),
        (0, [("", None, False)]),
        (0, [("def", "2563EB", True), (" ", None, False), ("save_log", "0F766E", False), ("(", "111827", False), ("message", "EA580C", False), ("):", "111827", False)]),
        (4, [("stamp", "111827", False), (" = ", "111827", False), ("time", "111827", False), (".", "111827", False), ("strftime", "0F766E", False), ("(", "111827", False), ('"%Y-%m-%d %H:%M:%S"', "16A34A", False), (")", "111827", False)]),
        (4, [("path", "111827", False), (" = ", "111827", False), ("Path", "0F766E", False), ("(", "111827", False), ('"logs.txt"', "16A34A", False), (")", "111827", False)]),
        (4, [("with", "8B5CF6", True), (" ", None, False), ("path", "111827", False), (".", "111827", False), ("open", "0F766E", False), ("(", "111827", False), ('"a"', "16A34A", False), (", encoding=", "111827", False), ('"utf-8"', "16A34A", False), (")", "111827", False), (" ", None, False), ("as", "8B5CF6", True), (" ", None, False), ("f", "111827", False), (":", "111827", False)]),
        (8, [("f", "111827", False), (".", "111827", False), ("write", "0F766E", False), ("(", "111827", False), ('f"[{stamp}] {message}\\n"', "16A34A", False), (")", "111827", False)]),
        (0, [("", None, False)]),
        (0, [("save_log", "0F766E", False), ("(", "111827", False), ('"첫 번째 예제 실행"', "16A34A", False), (")", "111827", False)]),
    ]

    for indent, parts in lines:
        p = document.add_paragraph(style="코드 블록")
        set_paragraph_shading(p, "F8FAFC")
        set_paragraph_border(p, color="CBD5E1", size=8, space=2)
        if parts == [("", None, False)]:
            add_run(p, " ", font="Consolas", size=9.5)
        else:
            add_code_line(p, indent, parts)


def add_table(document):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["요소", "권장 스타일", "설명"]
    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "DCE6F1")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, color="1A365D", size=10.5)

    rows = [
        ("장 제목", "Heading 1", "한 페이지에 한 번, 짧고 선명하게 작성"),
        ("절 제목", "Heading 2", "학습 단위를 빠르게 구분할 수 있도록 사용"),
        ("코드 블록", "코드 블록", "Consolas 9.5pt, 연한 배경, 줄 간격 1.1"),
        ("노트 박스", "강조 상자", "팁, 주의, 실무 메모를 분리해 전달"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            add_run(p, value, size=10.5)


def add_cover_page(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    add_run(p, "파이썬 중심 IT 도서 집필 템플릿", bold=True, color="1A365D", size=26)

    p2 = document.add_paragraph(style="Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "한국어 기술서를 위한 Word(.docx) 예시 템플릿", color="555555", size=13)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(18)
    add_run(p3, "권장 본문 폰트: 맑은 고딕 11pt", size=11)
    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p4, "권장 코드 폰트: Consolas 9.5pt", size=11)

    add_note_box(
        document,
        "사용 안내",
        "이 문서는 장 제목, 절 제목, 코드 예제, 표, 노트 박스, 경고 박스를 포함한 기본 집필 양식을 예시와 함께 보여줍니다.",
        "EEF5FF",
        "1D4ED8",
    )

    add_note_box(
        document,
        "편집 팁",
        "Word에서 제목 스타일을 수정하면 전체 문서의 장/절 디자인을 한 번에 바꿀 수 있습니다. 원고 작업 전 스타일 창에서 먼저 조정해 두면 편합니다.",
        "FFF7ED",
        "C2410C",
    )


def add_chapter_page(document):
    document.add_page_break()
    p = document.add_paragraph(style="Heading 1")
    add_run(p, "1장. 파이썬으로 시작하는 자동화", bold=True, color="1A365D", size=18)

    lead = document.add_paragraph(style="본문 강조")
    add_run(
        lead,
        "이 장은 독자가 파이썬의 기본 문법을 이해하고, 작은 자동화 스크립트를 직접 작성할 수 있도록 돕는 것을 목표로 합니다.",
        bold=True,
        color="1A365D",
        size=11,
    )

    body1 = document.add_paragraph()
    add_run(body1, "기술서는 설명의 흐름이 자연스러워야 합니다. 한 단락에는 하나의 메시지만 담고, 예제 코드는 바로 아래에 붙여 독자가 시선을 옮기지 않도록 구성하는 것이 좋습니다. 한국어 기술 문서는 영문 문서보다 문장 길이가 길어지는 경향이 있으므로, 문단 간격과 줄 간격을 조금 넉넉하게 두면 가독성이 좋아집니다.")

    body2 = document.add_paragraph()
    add_run(body2, "장 서두에는 학습 목표, 핵심 개념, 실습 결과물을 간결하게 배치해 독자가 이번 장에서 무엇을 얻는지 즉시 파악하도록 해주세요. 특히 입문서를 쓸 때는 용어 정의와 예제의 난이도 상승 폭을 세심하게 조절하는 것이 중요합니다.")

    add_note_box(
        document,
        "집필 메모",
        "장 시작 페이지에는 학습 목표 3개 이하, 핵심 키워드 5개 이하로 정리하면 독자가 부담 없이 진입할 수 있습니다.",
        "F0FDF4",
        "15803D",
    )

    p2 = document.add_paragraph(style="Heading 2")
    add_run(p2, "1.1 독자를 배려하는 설명 구조", bold=True, color="B85C38", size=14)

    body3 = document.add_paragraph()
    add_run(body3, "개념 설명 후 바로 짧은 예제를 제공하고, 그 다음에 예제의 실행 결과를 보여주는 3단 구성이 가장 안정적입니다. 이 패턴을 반복하면 독자는 책의 리듬을 빠르게 익히고, 저자는 장 전체의 일관성을 유지하기 쉬워집니다.")

    add_table(document)


def add_code_page(document):
    document.add_page_break()
    p = document.add_paragraph(style="Heading 1")
    add_run(p, "2장. 코드 예제 스타일 샘플", bold=True, color="1A365D", size=18)

    body = document.add_paragraph()
    add_run(body, "아래 코드는 한국어 기술서에서 자주 사용하는 짧은 파이썬 예제입니다. 키워드, 문자열, 함수 이름이 구분되도록 색을 나누었고, 배경색을 살짝 넣어 본문과 시각적으로 분리했습니다.")

    add_code_block(document)

    explain = document.add_paragraph()
    add_run(explain, "코드 블록 아래에는 반드시 실행 결과나 핵심 해설을 덧붙이는 것이 좋습니다. 독자는 코드를 보는 것만큼, 왜 이런 결과가 나오는지 이해하는 과정에서 더 많은 도움을 얻습니다.")

    add_note_box(
        document,
        "주의",
        "컬러 하이라이트는 인쇄 환경에서도 식별 가능해야 하므로, 색 차이만이 아니라 굵기와 여백 차이도 함께 사용하는 편이 안전합니다.",
        "FEF2F2",
        "B91C1C",
    )


def add_note_page(document):
    document.add_page_break()
    p = document.add_paragraph(style="Heading 1")
    add_run(p, "3장. 노트, 팁, 경고 박스 예시", bold=True, color="1A365D", size=18)

    body = document.add_paragraph()
    add_run(body, "기술 문서에서는 본문 흐름을 끊지 않으면서도 중요한 정보를 전달해야 하는 경우가 많습니다. 그럴 때 노트 박스를 사용하면, 독자가 핵심 포인트를 빠르게 스캔할 수 있습니다.")

    add_note_box(
        document,
        "TIP",
        "실습 전제 조건은 페이지 상단에 배치하고, 운영체제나 버전 차이 정보는 별도의 팁 박스로 묶어 주면 독자가 필요한 정보만 골라 읽기 쉽습니다.",
        "EFF6FF",
        "2563EB",
    )

    add_note_box(
        document,
        "실무 메모",
        "현업에서 바로 쓰는 예제는 파일 경로, 인코딩, 예외 처리처럼 실제 장애로 이어질 수 있는 포인트를 함께 다루는 편이 좋습니다.",
        "F0FDF4",
        "15803D",
    )

    add_note_box(
        document,
        "WARNING",
        "파괴적 명령어, 비용 발생 API 호출, 개인정보 처리 예제는 반드시 별도 경고 상자로 강조하고 안전한 대체 방법을 함께 안내하세요.",
        "FFF1F2",
        "BE123C",
    )

    body2 = document.add_paragraph()
    add_run(body2, "노트 박스는 지나치게 많으면 오히려 집중을 해치므로, 한 페이지 기준 1~3개 정도가 적당합니다. 설명이 길어진다면 박스보다 별도 소절로 분리하는 편이 더 읽기 좋습니다.")


def add_closing_page(document):
    document.add_page_break()
    p = document.add_paragraph(style="Heading 1")
    add_run(p, "4장. 원고 체크리스트", bold=True, color="1A365D", size=18)

    intro = document.add_paragraph()
    add_run(intro, "마지막 페이지에는 집필자가 반복적으로 확인해야 하는 항목을 체크리스트 형태로 두면 실제 작업 템플릿으로 활용하기 좋습니다.")

    checks = [
        "제목 체계가 Heading 1, Heading 2 스타일로 통일되어 있는가",
        "코드 예제에 실행 환경과 버전 정보가 포함되어 있는가",
        "이미지 또는 표 아래에 캡션과 해설이 있는가",
        "독자가 따라 할 수 있도록 단계 설명이 빠지지 않았는가",
        "용어 표기와 띄어쓰기가 문서 전반에서 일관적인가",
        "주의가 필요한 작업에 경고 박스를 사용했는가",
    ]
    for item in checks:
        p_item = document.add_paragraph()
        p_item.paragraph_format.left_indent = Cm(0.4)
        add_run(p_item, "□ ", bold=True, color="1A365D", size=12)
        add_run(p_item, item, size=11)

    add_note_box(
        document,
        "마무리",
        "이 파일은 바로 집필 시작이 가능하도록 기본 구성과 예시를 함께 담은 템플릿입니다. 필요하면 출판사 규격에 맞춰 여백, 본문 폰트, 제목 색상만 조정해 사용하면 됩니다.",
        "FFFBEB",
        "B45309",
    )


def main():
    document = Document()
    configure_page(document.sections[0])
    make_styles(document)
    add_header_footer(document.sections[0], "IT 도서 집필 템플릿")

    add_cover_page(document)
    add_chapter_page(document)
    add_code_page(document)
    add_note_page(document)
    add_closing_page(document)

    # Ensure all sections share the same page setup and header/footer style.
    for section in document.sections:
        configure_page(section)
        add_header_footer(section, "IT 도서 집필 템플릿")

    document.save(DOC_NAME)
    print(DOC_NAME)


if __name__ == "__main__":
    main()
